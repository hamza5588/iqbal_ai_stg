from flask import Blueprint, request, jsonify, session, send_file, after_this_request, render_template, current_app
from app.models.models import UserModel, LessonModel
from app.models.database_models import Lesson as DBLesson, LessonFAQ as DBLessonFAQ, LessonChatHistory as DBLessonChatHistory
from app.services.lesson_service import LessonService
from app.services.lesson.lesson_qa_graph import invoke_lesson_qa
from app.utils.decorators import login_required, teacher_required, student_required
from app.utils.db import get_db
from app.utils.groq_rate_limit import GroqRateLimitError, GroqBusyError
from werkzeug.datastructures import FileStorage
from werkzeug.exceptions import RequestEntityTooLarge
from sqlalchemy import func, or_, desc
import logging
import os
from io import BytesIO
import tempfile
import threading
import time
import re
from urllib.parse import urlparse

try:
    import redis
except Exception:  # pragma: no cover - optional dependency fallback
    redis = None

logger = logging.getLogger(__name__)

_lesson_qa_user_locks = {}
_lesson_qa_locks_lock = threading.Lock()
_lesson_qa_user_rate = {}
_lesson_qa_rate_lock = threading.Lock()
_lesson_qa_redis_client = None
_lesson_qa_redis_lock = threading.Lock()


def _normalize_faq_question_text(question: str) -> str:
    """Normalize semantically similar student questions into a stable FAQ key."""
    text = str(question or '').strip().lower()
    if not text:
        return ''

    # Expand common contractions to improve matching.
    contractions = {
        r"\bwhat's\b": "what is",
        r"\bwhats\b": "what is",
        r"\bwho's\b": "who is",
        r"\bwhere's\b": "where is",
        r"\bhow's\b": "how is",
        r"\bit's\b": "it is",
        r"\bcan't\b": "cannot",
        r"\bdon't\b": "do not",
    }
    for pattern, replacement in contractions.items():
        text = re.sub(pattern, replacement, text)

    text = re.sub(r"[^\w\s]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"^(what is\s+)?rephrased question\s+", "", text).strip()

    # Remove polite wrappers and request phrasing.
    text = re.sub(r"^(please\s+)?(can|could|would)\s+you\s+(tell|explain|define|describe)\s+", "", text)
    text = re.sub(r"^(can|could)\s+you\s+tell\s+me\s+", "", text)
    text = re.sub(r"^please\s+", "", text).strip()

    # Collapse semantic variants such as:
    # "what is a magnet", "what is the meaning of a magnet", "meaning of magnet".
    core = text
    for prefix in (
        "what is the meaning of ",
        "what is meaning of ",
        "what are the different types of ",
        "meaning of ",
        "define ",
        "explain ",
        "what is ",
        "what are ",
    ):
        if core.startswith(prefix):
            core = core[len(prefix):].strip()
            break

    core = re.sub(r"^(a|an|the)\s+", "", core).strip()
    core = re.sub(r"\s+", " ", core)
    if not core:
        return text
    return f"what is {core}"


def _faq_display_question(question_key: str) -> str:
    """Convert normalized key into a readable FAQ question."""
    cleaned = str(question_key or "").strip()
    if not cleaned:
        return ""
    if not cleaned.endswith("?"):
        cleaned += "?"
    return cleaned[:1].upper() + cleaned[1:]


def _get_lesson_qa_user_lock(user_id: int) -> threading.Lock:
    with _lesson_qa_locks_lock:
        if user_id not in _lesson_qa_user_locks:
            _lesson_qa_user_locks[user_id] = threading.Lock()
        return _lesson_qa_user_locks[user_id]


def _get_lesson_qa_redis_client():
    """
    Best-effort Redis client for cross-worker lesson QA rate limiting.
    Falls back to in-memory limiter if Redis is unavailable.
    """
    global _lesson_qa_redis_client
    if _lesson_qa_redis_client is not None:
        return _lesson_qa_redis_client
    if redis is None:
        return None

    with _lesson_qa_redis_lock:
        if _lesson_qa_redis_client is not None:
            return _lesson_qa_redis_client
        redis_url = (
            os.getenv("LESSON_QA_RATE_LIMIT_REDIS_URL")
            or os.getenv("CELERY_BROKER_URL")
            or "redis://localhost:6379/0"
        )
        try:
            parsed = urlparse(redis_url)
            if parsed.scheme.startswith("redis"):
                _lesson_qa_redis_client = redis.Redis.from_url(
                    redis_url,
                    socket_connect_timeout=1.0,
                    socket_timeout=1.0,
                    health_check_interval=30,
                )
                _lesson_qa_redis_client.ping()
                logger.info("Lesson QA rate limiter: using Redis backend at %s", redis_url)
            else:
                logger.warning(
                    "Lesson QA rate limiter: non-redis broker URL '%s'; falling back to in-memory.",
                    redis_url,
                )
        except Exception as exc:
            logger.warning(
                "Lesson QA rate limiter: Redis unavailable (%s); using in-memory fallback.",
                exc,
            )
            _lesson_qa_redis_client = None
    return _lesson_qa_redis_client


def _check_and_record_lesson_qa_rate(user_id: int) -> tuple[bool, int]:
    """Per-user burst throttling for lesson Q&A requests."""
    max_requests = int(os.getenv("LESSON_QA_USER_RATE_LIMIT_COUNT", "6"))
    window_seconds = int(os.getenv("LESSON_QA_USER_RATE_LIMIT_WINDOW_SECONDS", "10"))
    now = time.time()

    # Preferred path: Redis sorted-set sliding window (shared across workers).
    redis_client = _get_lesson_qa_redis_client()
    if redis_client is not None:
        key = f"lesson_qa_rate:user:{int(user_id)}"
        cutoff = now - window_seconds
        now_ms = int(now * 1000)
        unique_member = f"{now_ms}:{threading.get_ident()}"
        try:
            pipe = redis_client.pipeline()
            pipe.zremrangebyscore(key, 0, cutoff)
            pipe.zcard(key)
            _, current_count = pipe.execute()
            if int(current_count or 0) >= max_requests:
                oldest_entries = redis_client.zrange(key, 0, 0, withscores=True)
                if oldest_entries:
                    oldest_score = float(oldest_entries[0][1])
                    retry_after = max(1, int(window_seconds - (now - oldest_score)))
                else:
                    retry_after = window_seconds
                return False, retry_after

            pipe = redis_client.pipeline()
            pipe.zadd(key, {unique_member: now})
            pipe.expire(key, max(window_seconds + 2, 2 * window_seconds))
            pipe.execute()
            return True, 0
        except Exception as exc:
            # Safe fallback under Redis outages.
            logger.warning(
                "Lesson QA rate limiter Redis error (%s); falling back to in-memory for this request.",
                exc,
            )

    # Fallback path: in-memory per-process limiter.
    with _lesson_qa_rate_lock:
        history = _lesson_qa_user_rate.get(user_id, [])
        cutoff = now - window_seconds
        history = [ts for ts in history if ts >= cutoff]
        if len(history) >= max_requests:
            oldest = history[0]
            retry_after = max(1, int(window_seconds - (now - oldest)))
            _lesson_qa_user_rate[user_id] = history
            return False, retry_after
        history.append(now)
        _lesson_qa_user_rate[user_id] = history

    return True, 0


def _canonicalize_and_log_in_background(
    app,
    lesson_id: int,
    user_id: int,
    question: str,
    chat_history_id: int,
    api_key: str,
) -> None:
    """Canonicalize and FAQ-log question without blocking the API response."""
    with app.app_context():
        from app.models.models import LessonChatHistory, LessonFAQ

        canonical = question
        try:
            service = LessonService(api_key=api_key)
            canonical = service.canonicalize_question(int(lesson_id), question) or question
        except Exception as e:
            logger.error(f"Background canonicalization failed for lesson {lesson_id}, user {user_id}: {str(e)}")
            canonical = question
        canonical = _normalize_faq_question_text(canonical or question) or (question or '').strip()

        try:
            LessonChatHistory.update_canonical_question(chat_history_id, canonical)
        except Exception as e:
            logger.error(
                f"Failed to update canonical_question for lesson_chat_history {chat_history_id}: {str(e)}"
            )

        try:
            LessonFAQ.log_question(int(lesson_id), canonical)
            logger.info(f"Question logged to FAQ table for lesson {lesson_id}: {canonical}")
        except Exception as e:
            logger.error(f"Error logging question to FAQ table: {str(e)}")


def _get_lesson_api_key() -> str:
    """
    Resolve the API key for lesson generation/download.
    Priority:
    1) Session `groq_api_key`
    2) `GROQ_API_KEY` environment variable
    3) `OPENAI_API_KEY` environment variable
    """
    return (
        session.get('groq_api_key')
        or os.getenv('GROQ_API_KEY')
        or os.getenv('OPENAI_API_KEY')
        or ''
    )


# Create a custom logger for lesson checks
lesson_check_logger = logging.getLogger('lesson_check')
lesson_check_logger.setLevel(logging.DEBUG)

# Ensure logs directory exists
os.makedirs('logs', exist_ok=True)

lesson_check_handler = logging.FileHandler('logs/lesson_check.log')
lesson_check_handler.setLevel(logging.DEBUG)
lesson_check_formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
lesson_check_handler.setFormatter(lesson_check_formatter)
lesson_check_logger.addHandler(lesson_check_handler)
lesson_check_logger.info("=" * 80)
lesson_check_logger.info("Lesson Check Logger Initialized")
lesson_check_logger.info("=" * 80)

bp = Blueprint('lesson_routes', __name__)

@bp.route('/create_lesson', methods=['POST'])
# @teacher_required
def create_lesson():
    """Create a new lesson or answer a question based on user input"""
    logger.info("=== CREATE_LESSON REQUEST RECEIVED ===")
    logger.info(f"Request method: {request.method}")
    logger.info(f"Request path: {request.path}")
    logger.info(f"Content-Type: {request.content_type}")
    logger.info(f"Content-Length: {request.content_length if request.content_length else 'Unknown'}")
    logger.info(f"Has files: {'file' in request.files}")
    if 'file' in request.files:
        file = request.files['file']
        logger.info(f"File in request: {file.filename if file.filename else 'No filename'}")
    
    try:
        # Get form data for lesson configuration
        lesson_title = request.form.get('lessonTitle', '')
        lesson_prompt = request.form.get('lessonPrompt', '')
        focus_area = request.form.get('focusArea', '')
        grade_level = request.form.get('gradeLevel', '')
        additional_notes = request.form.get('additionalNotes', '')
        
        logger.info(f"Form data - Title: {lesson_title}, Focus: {focus_area}, Grade: {grade_level}")
        
        # Get API key (session or environment fallback)
        api_key = _get_lesson_api_key()
        if not api_key:
            logger.warning("API key not found in session")
            return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
        
        # Initialize lesson service
        logger.info("Initializing lesson service...")
        lesson_service = LessonService(api_key=api_key)
        
        # Check if file is provided - if yes, always treat as lesson generation
        if 'file' in request.files and request.files['file'].filename != '':
            # User uploaded a file - proceed with lesson generation regardless of prompt type
            logger.info("File uploaded - proceeding with lesson generation")
            file = request.files['file']
            
            # Check file size before processing
            if file and file.filename:
                # Get file size
                file.seek(0, os.SEEK_END)
                file_size = file.tell()
                file.seek(0)  # Reset file pointer
                
                logger.info(f"File upload attempt: {file.filename}, size: {file_size / (1024*1024):.2f}MB")
                
                if file_size > 100 * 1024 * 1024:  # 100MB
                    return jsonify({'error': 'File size exceeds 100MB limit'}), 400
            
            # Validate required fields for lesson generation (prompt is now optional since we use interactive chat)
            if not lesson_title or not focus_area or not grade_level:
                return jsonify({'error': 'All required fields must be filled for lesson generation'}), 400
            
            # Check if lesson title already exists for this teacher
            if LessonModel.check_title_exists(session['user_id'], lesson_title):
                return jsonify({'error': 'This lesson title is already used. Please choose a different title.'}), 400
            
            # Check file type
            allowed_extensions = {'.pdf', '.doc', '.docx', '.txt'}
            file_ext = os.path.splitext(file.filename.lower())[1]
            if file_ext not in allowed_extensions:
                return jsonify({'error': 'File type not supported. Please upload PDF, DOC, DOCX, or TXT files.'}), 400
            
            # Get extraction toggle values
            table_extraction = request.form.get('tableExtraction', 'false').lower() == 'true'
            image_extraction = request.form.get('imageExtraction', 'false').lower() == 'true'
            

            # Prepare lesson details
            lesson_details = {
                'lesson_title': lesson_title,
                'lesson_prompt': lesson_prompt,
                'focus_area': focus_area,
                'grade_level': grade_level,
                'additional_notes': additional_notes,
                'table_extraction': table_extraction,
                'image_extraction': image_extraction
            }
            
            # Keep a copy of file bytes for RAG ingest (PDF only) after lesson is created
            file.seek(0)
            file_bytes_for_rag = file.read()
            file.seek(0)
            
            # Process the file first to create vector store
            process_result = lesson_service.process_file(file, lesson_details)
            

            if 'error' in process_result:
                return jsonify({'error': process_result['error']}), 400
            
            # Create a draft lesson entry in database
            # Use a default summary since prompt is no longer required
            summary_text = f"Lesson on {focus_area} for {grade_level} grade students"
            if additional_notes:
                summary_text = additional_notes[:200] + "..." if len(additional_notes) > 200 else additional_notes
            
            lesson_id = LessonModel.create_lesson(
                teacher_id=session['user_id'],
                title=lesson_title,
                summary=summary_text,
                learning_objectives='',
                content='',  # Will be filled when lesson is complete
                grade_level=grade_level,
                focus_area=focus_area,
                is_public=True  # Make lessons public by default
            )
            
            if not lesson_id:
                return jsonify({'error': 'Failed to save lesson to database'}), 500
            
            # Store RAG thread id when lesson is first saved (PDF only); then ingest for "Ask Question"
            teacher_id = session['user_id']
            if file_ext == '.pdf':
                rag_thread_id = f"user_{teacher_id}_lesson_{lesson_id}"
                lesson_model = LessonModel(lesson_id)
                lesson_model.update_lesson(rag_thread_id=rag_thread_id)
                logger.info("Lesson %s linked to RAG thread %s for student Q&A", lesson_id, rag_thread_id)
                if file_bytes_for_rag:
                    try:
                        from app.utils.rag_service import ingest_pdf
                        ingest_result = ingest_pdf(
                            file_bytes_for_rag,
                            rag_thread_id,
                            filename=file.filename,
                            user_id=teacher_id,
                        )
                        if ingest_result.get('error'):
                            logger.warning("Lesson RAG ingest failed (non-fatal): %s", ingest_result.get('error'))
                    except Exception as rag_e:
                        logger.warning("Lesson RAG ingest failed (non-fatal): %s", rag_e)
            
            # Get the greeting message from process_result (no LLM call was made)
            greeting_message = process_result.get('lesson', 'Your file has been uploaded successfully.')
            
            # Get the lesson response
            lesson_response = LessonModel.get_lesson_by_id(lesson_id)
            lesson_response['title'] = lesson_title
            lesson_response['grade_level'] = grade_level
            lesson_response['focus_area'] = focus_area
            lesson_response['isFinalized'] = False
            
            # Return greeting message without calling LLM
            # LLM will only be called when user sends a query via interactive_chat endpoint
            images_processing = process_result.get('images_processing', False)
            return jsonify({
                'success': True,
                'response_type': 'file_uploaded',
                'lesson_id': lesson_id,
                'lesson': lesson_response,
                'ai_response': greeting_message,
                'complete_lesson': 'no',
                'message': 'File uploaded successfully! You can now start chatting to create your lesson.',
                'file_processed': True,
                'filename': process_result.get('filename_processed', file.filename),
                'images_processing': images_processing  # Flag to indicate images are processing in background
            })   
        
        else:
            # No file uploaded - use additional_notes or lesson_prompt if provided, otherwise require file
            query_text = additional_notes or lesson_prompt or ''
            
            if query_text:
                # Analyze query to determine if it's a question or lesson generation request
                query_analysis = lesson_service.analyze_user_query(query_text)
                logger.info(f"Query analysis: {query_analysis}")
                
                if query_analysis['query_type'] == 'QUESTION':
                    # User is asking a question - answer it using available lessons
                    logger.info("User query detected as question - answering using available lessons")
                    
                    # Get available lesson IDs for context
                    available_lessons = lesson_service._get_available_lesson_ids()
                    
                    # Answer the question using available lesson content
                    answer_result = lesson_service.answer_general_question(query_text, available_lessons)
                    
                    return jsonify({
                        'success': True,
                        'response_type': 'question_answer',
                        'answer': answer_result['answer'],
                        'source': answer_result.get('source', 'unknown'),
                        'confidence': answer_result.get('confidence', 0.0),
                        'lesson_context': answer_result.get('lesson_context', ''),
                        'relevant_lessons': answer_result.get('relevant_lessons', []),
                        'query_analysis': query_analysis,
                        'message': 'Question answered based on available lesson content'
                    })
            
            # User wants to generate a lesson but no file provided
            return jsonify({'error': 'File is required for lesson generation. Please upload a PDF, DOC, DOCX, or TXT file.'}), 400
        
    except RequestEntityTooLarge:
        logger.error("=== CREATE_LESSON ERROR: File upload rejected - File too large ===")
        return jsonify({'error': 'File size exceeds maximum allowed size (100MB)'}), 413
    except Exception as e:
        logger.error("=== CREATE_LESSON ERROR ===")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error message: {str(e)}")
        logger.error(f"Request path: {request.path}")
        logger.error(f"Request method: {request.method}")
        logger.error(f"Content-Length: {request.content_length if request.content_length else 'Unknown'}")
        logger.error(f"Exception details:", exc_info=True)
        return jsonify({'error': f'Failed to process request: {str(e)}'}), 500

@bp.route('/create', methods=['POST'])
@login_required
def create_lesson_simple():
    """Create a lesson from JSON data (for finalized lessons from chat)"""
    try:
        if 'user_id' not in session:
            return jsonify({'error': 'Not authenticated'}), 401
        
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        title = data.get('title', '').strip()
        content = data.get('content', '').strip()
        focus_area = data.get('focus_area', 'General')
        grade_level = data.get('grade_level', 'General')
        summary = data.get('summary', '') or data.get('additional_notes', '')
        rag_thread_id = (data.get('rag_thread_id') or data.get('thread_id') or '').strip() or None
        
        if not title:
            return jsonify({'error': 'Title is required'}), 400
        
        if not content:
            return jsonify({'error': 'Content is required'}), 400
        
        # Check if lesson title already exists for this teacher
        if LessonModel.check_title_exists(session['user_id'], title):
            return jsonify({'error': 'This lesson title is already used. Please choose a different title.'}), 400
        
        # Create the lesson
        lesson_id = LessonModel.create_lesson(
            teacher_id=session['user_id'],
            title=title,
            summary=summary or f"Lesson on {focus_area} for {grade_level}",
            learning_objectives='',
            content=content,
            grade_level=grade_level,
            focus_area=focus_area,
            is_public=True,
            status='finalized',
            rag_thread_id=rag_thread_id
        )
        
        if not lesson_id:
            return jsonify({'error': 'Failed to save lesson to database'}), 500
        
        # Get the created lesson
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        
        return jsonify({
            'success': True,
            'lesson': lesson,
            'id': lesson_id,
            'message': 'Lesson saved successfully'
        })
        
    except Exception as e:
        logger.error(f"Error creating lesson: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to create lesson: {str(e)}'}), 500

@bp.route('/ask_question_general', methods=['POST'])
@login_required
def ask_general_question():
    """Ask a general question that will be answered using available lesson content"""
    try:
        data = request.get_json()
        question = data.get('question', '').strip()
        
        if not question:
            return jsonify({'error': 'Question is required'}), 400
        
        # Get API key from session
        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
        
        # Initialize lesson service
        lesson_service = LessonService(api_key=api_key)
        
        # Get conversation history from all lessons for context
        from app.models.models import LessonChatHistory
        user_id = session['user_id']
        
        # Get recent conversation history from all lessons
        conversation_history = []
        available_lessons = lesson_service._get_available_lesson_ids()
        
        # Get recent Q&A from all lessons (last 5 from each lesson)
        for lesson_id in available_lessons[:5]:  # Limit to first 5 lessons to avoid too much context
            try:
                lesson_history = LessonChatHistory.get_lesson_chat_history(lesson_id, user_id)
                conversation_history.extend(lesson_history[-2:])  # Last 2 from each lesson
            except Exception as e:
                logger.warning(f"Error getting history for lesson {lesson_id}: {str(e)}")
                continue
        
        # Sort by timestamp and take most recent
        conversation_history = sorted(conversation_history, key=lambda x: x.get('created_at', ''), reverse=True)[:10]
        
        # Analyze the user's query to determine intent
        query_analysis = lesson_service.analyze_user_query(question)
        logger.info(f"General query analysis: {query_analysis}")
        
        if not available_lessons:
            return jsonify({
                'success': True,
                'response_type': 'no_content',
                'answer': 'I don\'t have any lesson content available to answer your question. Please upload some educational materials first.',
                'source': 'no_content',
                'confidence': 0.0,
                'query_analysis': query_analysis,
                'message': 'No lesson content available'
            })
        
        # Answer the question using available lesson content with conversation context
        answer_result = lesson_service.answer_general_question(question, available_lessons, conversation_history)
        
        return jsonify({
            'success': True,
            'response_type': 'question_answer',
            'answer': answer_result['answer'],
            'source': answer_result.get('source', 'unknown'),
            'confidence': answer_result.get('confidence', 0.0),
            'lesson_context': answer_result.get('lesson_context', ''),
            'relevant_lessons': answer_result.get('relevant_lessons', []),
            'query_analysis': query_analysis,
            'message': 'Question answered based on available lesson content'
        })
        
    except Exception as e:
        logger.error(f"Error answering general question: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to answer question: {str(e)}'}), 500

@bp.route('/my_lessons', methods=['GET'])
@teacher_required
def get_my_lessons():
    """Get teacher lessons with server-side pagination."""
    try:
        page = request.args.get('page', default=1, type=int)
        per_page = request.args.get('per_page', default=10, type=int)
        search_term = (request.args.get('q') or '').strip() or None
        # latest_only=true (default) hides older versions so only the latest version
        # of each logical lesson is returned — avoids duplicate rows in the UI table.
        latest_only = request.args.get('latest_only', 'true').lower() != 'false'
        result = LessonModel.get_lessons_by_teacher_paginated(
            teacher_id=session['user_id'],
            page=page,
            per_page=per_page,
            search_term=search_term,
            latest_only=latest_only,
        )
        return jsonify({
            'success': True,
            'lessons': result['lessons'],
            'page': result['page'],
            'per_page': result['per_page'],
            'total': result['total'],
            'total_pages': result['total_pages'],
        })
    except Exception as e:
        logger.error(f"Error getting teacher lessons: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to get lessons: {str(e)}'}), 500

@bp.route('/browse_lessons', methods=['GET'])
@student_required
def browse_lessons():
    """Browse public lessons for students with server-side pagination."""
    try:
        grade_level = request.args.get('grade_level')
        focus_area = request.args.get('focus_area')
        page = request.args.get('page', default=1, type=int)
        per_page = request.args.get('per_page', default=10, type=int)
        search_term = (request.args.get('q') or '').strip() or None

        result = LessonModel.get_public_latest_lessons_paginated(
            grade_level=grade_level,
            focus_area=focus_area,
            page=page,
            per_page=per_page,
            search_term=search_term,
        )
        return jsonify({
            'success': True,
            'lessons': result['lessons'],
            'page': result['page'],
            'per_page': result['per_page'],
            'total': result['total'],
            'total_pages': result['total_pages'],
        })
    except Exception as e:
        logger.error(f"Error browsing lessons: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to browse lessons: {str(e)}'}), 500

@bp.route('/search_lessons', methods=['GET'])
@login_required
def search_lessons():
    """Search lessons for both teachers and students"""
    try:
        search_term = request.args.get('q', '')
        grade_level = request.args.get('grade_level')
        
        if not search_term:
            return jsonify({'error': 'Search term is required'}), 400
        
        lessons = LessonModel.search_lessons(search_term, grade_level=grade_level)
        return jsonify({
            'success': True,
            'lessons': lessons
        })
    except Exception as e:
        logger.error(f"Error searching lessons: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to search lessons: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>', methods=['GET'])
@login_required
def get_lesson(lesson_id):
    """Get a specific lesson by ID"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        # Check if user can access this lesson
        user_id = session.get('user_id')
        user_role = session.get('role', 'student')
        lesson_teacher_id = lesson.get('teacher_id')
        is_public = lesson.get('is_public', False)
        
        # Users can access their own lessons (regardless of role)
        if lesson_teacher_id == user_id:
            pass  # Allow access - user owns this lesson
        # If lesson is public, allow access for anyone
        elif is_public:
            pass  # Allow access - lesson is public
        # Otherwise deny access
        else:
            logger.info(f"Access denied for user {user_id} (role: {user_role}) to lesson {lesson_id} (teacher: {lesson_teacher_id}, public: {is_public})")
            return jsonify({'error': 'Access denied'}), 403
        
        return jsonify({
            'success': True,
            'lesson': lesson
        })
    except Exception as e:
        logger.error(f"Error getting lesson: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to get lesson: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/view', methods=['GET'])
@login_required
def view_lesson(lesson_id):
    """Get a lesson with its versions for viewing"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        # Check if user can access this lesson
        user_role = session.get('role', 'student')
        user_id = session.get('user_id')
        lesson_teacher_id = lesson.get('teacher_id')
        is_public = lesson.get('is_public', False)
        
        # Users can access their own lessons (regardless of role)
        if lesson_teacher_id == user_id:
            pass  # Allow access - user owns this lesson
        # If lesson is public, allow access for anyone
        elif is_public:
            pass  # Allow access - lesson is public
        # Otherwise deny access
        else:
            logger.info(f"Access denied for user {user_id} (role: {user_role}) to lesson {lesson_id} (teacher: {lesson_teacher_id}, public: {is_public})")
            return jsonify({'error': 'Access denied'}), 403
        
        # get_lesson_versions queries by (id == root_id OR parent_lesson_id == root_id).
        # All child versions store parent_lesson_id pointing to the ROOT lesson, so we
        # must resolve to the root before querying — otherwise only the clicked version
        # is returned (its own row) with no history.
        root_lesson_id = lesson.get('parent_lesson_id') or lesson_id
        versions = LessonModel.get_lesson_versions(root_lesson_id)

        return jsonify({
            'success': True,
            'lesson': lesson,
            'versions': versions
        })
    except Exception as e:
        logger.error(f"Error viewing lesson: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to view lesson: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>', methods=['PUT'])
@teacher_required
def update_lesson(lesson_id):
    """Update a lesson (teacher only, and only their own lessons)"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        # Check if the lesson belongs to the current teacher
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        data = request.get_json()
        
        # Check if new title already exists for this teacher (if title is being changed)
        new_title = data.get('title')
        if new_title and new_title != lesson['title'] and LessonModel.check_title_exists(session['user_id'], new_title, exclude_lesson_id=lesson_id):
            return jsonify({'error': 'This lesson title is already used. Please choose a different title.'}), 400
        
        lesson_model = LessonModel(lesson_id)
        
        # Update lesson - if content is provided, also update original_content if it's empty
        content_to_save = data.get('content')
        if content_to_save:
            # Check if original_content is empty, if so, set it to the content being saved
            lesson = LessonModel.get_lesson_by_id(lesson_id)
            if lesson and (not lesson.get('original_content') or lesson.get('original_content').strip() == ''):
                # Update original_content as well if it's empty using ORM
                db = get_db()
                lesson_obj = db.query(DBLesson).filter(DBLesson.id == lesson_id).first()
                if lesson_obj:
                    lesson_obj.original_content = content_to_save
                    db.commit()
                    logger.info(f"Updated original_content for lesson {lesson_id} since it was empty")
        
        success = lesson_model.update_lesson(
            title=data.get('title'),
            summary=data.get('summary'),
            learning_objectives=data.get('learning_objectives'),
            focus_area=data.get('focus_area'),
            grade_level=data.get('grade_level'),
            content=content_to_save,
            is_public=data.get('is_public')
        )
        
        if success:
            return jsonify({
                'success': True,
                'message': 'Lesson updated successfully'
            })
        else:
            return jsonify({'error': 'Failed to update lesson'}), 500
            
    except Exception as e:
        logger.error(f"Error updating lesson: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to update lesson: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>', methods=['DELETE'])
@teacher_required
def delete_lesson(lesson_id):
    """Delete a lesson (teacher only, and only their own lessons)"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        # Check if the lesson belongs to the current teacher
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        lesson_model = LessonModel(lesson_id)
        success = lesson_model.delete_lesson()
        
        if success:
            return jsonify({
                'success': True,
                'message': 'Lesson deleted successfully'
            })
        else:
            return jsonify({'error': 'Failed to delete lesson'}), 500
            
    except Exception as e:
        logger.error(f"Error deleting lesson: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to delete lesson: {str(e)}'}), 500

@bp.route('/download_lesson/<int:lesson_id>', methods=['GET'])
@login_required
def download_lesson(lesson_id):
    """Download a lesson as DOCX file"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        # Check if user can access this lesson
        user_role = session.get('role', 'student')
        user_id = session.get('user_id')
        lesson_teacher_id = lesson.get('teacher_id')
        is_public = lesson.get('is_public', False)
        
        # Users can access their own lessons (regardless of role)
        if lesson_teacher_id == user_id:
            pass  # Allow access - user owns this lesson
        # If lesson is public, allow access for anyone
        elif is_public:
            pass  # Allow access - lesson is public
        # Otherwise deny access
        else:
            logger.info(f"Access denied for user {user_id} (role: {user_role}) to lesson {lesson_id} (teacher: {lesson_teacher_id}, public: {is_public})")
            return jsonify({'error': 'Access denied'}), 403
        
        # Create lesson structure for DOCX generation
        lesson_data = {
            'title': lesson['title'],
            'summary': lesson['summary'] or '',
            'learning_objectives': [lesson['learning_objectives']] if lesson['learning_objectives'] else [],
            'sections': [{'heading': 'Lesson Content', 'content': lesson['content']}],
            'key_concepts': [],
            'activities': [],
            'quiz': []
        }

        # Generate DOCX (no API key required for download)
        lesson_service = LessonService()
        docx_bytes = lesson_service._create_docx(lesson_data)
        
        # Create filename
        filename = lesson['title'].replace(' ', '_') + '.docx'
        
        # Create BytesIO object
        docx_buffer = BytesIO(docx_bytes)
        docx_buffer.seek(0)
        
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Download error: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to download lesson: {str(e)}'}), 500 

@bp.route('/download_lesson_ppt/<int:lesson_id>', methods=['GET'])
@login_required
def download_lesson_ppt(lesson_id):
    """Download lesson as PowerPoint presentation"""
    lesson = LessonModel.get_lesson_by_id(lesson_id)
    if not lesson:
        return jsonify({'error': 'Lesson not found'}), 404

    # Check if user can access this lesson
    user_role = session.get('role', 'student')
    user_id = session.get('user_id')
    lesson_teacher_id = lesson.get('teacher_id')
    is_public = lesson.get('is_public', False)
    
    # Users can access their own lessons (regardless of role)
    if lesson_teacher_id == user_id:
        pass  # Allow access - user owns this lesson
    # If lesson is public, allow access for anyone
    elif is_public:
        pass  # Allow access - lesson is public
    # Otherwise deny access
    else:
        logger.info(f"Access denied for user {user_id} (role: {user_role}) to lesson {lesson_id} (teacher: {lesson_teacher_id}, public: {is_public})")
        return jsonify({'error': 'Access denied'}), 403

    # Prepare lesson data for PPT generation
    lesson_data = {
        'title': lesson['title'],
        'summary': lesson['summary'] or '',
        'learning_objectives': [lesson['learning_objectives']] if lesson['learning_objectives'] else [],
        'sections': [{'heading': 'Lesson Content', 'content': lesson['content']}],
        'key_concepts': [],
        'activities': [],
        'quiz': []
    }
    
    try:
        # Generate PPT (no API key required for download)
        lesson_service = LessonService()
        ppt_bytes = lesson_service.create_ppt(lesson_data)
        
        # Create filename
        filename = lesson['title'].replace(' ', '_') + '.pptx'
        
        # Create response
        from io import BytesIO
        ppt_buffer = BytesIO(ppt_bytes)
        ppt_buffer.seek(0)
        
        # Delete FAISS index after successful download (best-effort cleanup)
        try:
            lesson_service._delete_faiss_index(lesson_id)
            logger.info(f"Deleted FAISS index after PPT download for lesson {lesson_id}")
        except Exception as e:
            logger.warning(f"Failed to delete FAISS index for lesson {lesson_id}: {str(e)}")
        
        return send_file(
            ppt_buffer,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.presentationml.presentation'
        )
    except Exception as e:
        logger.error(f"PPT generation error: {str(e)}")
        return jsonify({'error': f'Failed to generate PPT: {str(e)}'}), 500

@bp.route('/download_lesson_pdf/<int:lesson_id>', methods=['GET'])
@login_required
def download_lesson_pdf(lesson_id):
    """Download a lesson as PDF file"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404

        # Check if user can access this lesson
        user_role = session.get('role', 'student')
        user_id = session.get('user_id')
        lesson_teacher_id = lesson.get('teacher_id')
        is_public = lesson.get('is_public', False)
        
        # Users can access their own lessons (regardless of role)
        if lesson_teacher_id == user_id:
            pass  # Allow access - user owns this lesson
        # If lesson is public, allow access for anyone
        elif is_public:
            pass  # Allow access - lesson is public
        # Otherwise deny access
        else:
            logger.info(f"Access denied for user {user_id} (role: {user_role}) to lesson {lesson_id} (teacher: {lesson_teacher_id}, public: {is_public})")
            return jsonify({'error': 'Access denied'}), 403

        # Try to generate PDF using reportlab (preferred method)
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
            
            # Create PDF buffer
            pdf_buffer = BytesIO()
            
            # Create PDF document
            doc = SimpleDocTemplate(pdf_buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # Define custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=TA_CENTER
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12,
                spaceBefore=20
            )
            
            content_style = ParagraphStyle(
                'CustomContent',
                parent=styles['Normal'],
                fontSize=11,
                spaceAfter=12,
                alignment=TA_LEFT
            )
            
            # Build PDF content
            story = []
            
            # Add title
            story.append(Paragraph(lesson['title'], title_style))
            story.append(Spacer(1, 0.2*inch))
            
            # Add summary if available
            if lesson.get('summary'):
                story.append(Paragraph("Summary", heading_style))
                story.append(Paragraph(lesson['summary'], content_style))
                story.append(Spacer(1, 0.1*inch))
            
            # Add learning objectives if available
            if lesson.get('learning_objectives'):
                story.append(Paragraph("Learning Objectives", heading_style))
                story.append(Paragraph(lesson['learning_objectives'], content_style))
                story.append(Spacer(1, 0.1*inch))
            
            # Add focus area and grade level
            if lesson.get('focus_area'):
                story.append(Paragraph("Focus Area", heading_style))
                story.append(Paragraph(lesson['focus_area'], content_style))
                story.append(Spacer(1, 0.1*inch))
            
            if lesson.get('grade_level'):
                story.append(Paragraph("Grade Level", heading_style))
                story.append(Paragraph(lesson['grade_level'], content_style))
                story.append(Spacer(1, 0.1*inch))
            
            # Add main content
            if lesson.get('content'):
                story.append(Paragraph("Lesson Content", heading_style))
                # Split content by paragraphs and create Paragraph objects
                content_paragraphs = lesson['content'].split('\n\n')
                for para in content_paragraphs:
                    if para.strip():
                        story.append(Paragraph(para.strip(), content_style))
                        story.append(Spacer(1, 0.05*inch))
            
            # Build PDF
            doc.build(story)
            pdf_buffer.seek(0)
            
            # Create filename
            filename = lesson['title'].replace(' ', '_').replace('/', '_') + '.pdf'
            
            # Delete FAISS index after successful download
            try:
                lesson_service = LessonService(api_key=api_key)
                lesson_service._delete_faiss_index(lesson_id)
                logger.info(f"Deleted FAISS index after PDF download for lesson {lesson_id}")
            except Exception as e:
                logger.warning(f"Failed to delete FAISS index for lesson {lesson_id}: {str(e)}")
            
            return send_file(
                pdf_buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/pdf'
            )
            
        except ImportError:
            logger.warning("ReportLab not available, falling back to LibreOffice method")
            # Fallback to LibreOffice method if reportlab is not available
            pass

        # Fallback: Generate DOCX first and convert with LibreOffice
        lesson_data = {
            'title': lesson['title'],
            'summary': lesson['summary'] or '',
            'learning_objectives': [lesson['learning_objectives']] if lesson['learning_objectives'] else [],
            'sections': [{'heading': 'Lesson Content', 'content': lesson['content']}],
            'key_concepts': [],
            'activities': [],
            'quiz': []
        }
        
        # Get API key from session
        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
        
        lesson_service = LessonService(api_key=api_key)
        docx_bytes = lesson_service._create_docx(lesson_data)

        # Save DOCX to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as docx_file:
            docx_file.write(docx_bytes)
            docx_path = docx_file.name

        pdf_path = os.path.splitext(docx_path)[0] + '.pdf'
        try:
            import subprocess
            # Use LibreOffice to convert DOCX to PDF
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path
            ], capture_output=True, timeout=30)
            
            if result.returncode != 0:
                raise Exception(f'LibreOffice conversion failed: {result.stderr.decode()}')

            @after_this_request
            def cleanup(response):
                try:
                    os.remove(docx_path)
                except Exception:
                    pass
                try:
                    os.remove(pdf_path)
                except Exception:
                    pass
                return response

            # Delete FAISS index after successful download
            try:
                lesson_service = LessonService(api_key=api_key)
                lesson_service._delete_faiss_index(lesson_id)
                logger.info(f"Deleted FAISS index after PDF download for lesson {lesson_id}")
            except Exception as e:
                logger.warning(f"Failed to delete FAISS index for lesson {lesson_id}: {str(e)}")
            
            return send_file(pdf_path, as_attachment=True, download_name=lesson['title'].replace('/', '_') + '.pdf', mimetype='application/pdf')
            
        except Exception as e:
            # Clean up files on error
            if os.path.exists(docx_path):
                os.remove(docx_path)
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
            logger.error(f"LibreOffice PDF generation error: {str(e)}")
            return jsonify({'error': 'PDF generation failed. Please try downloading as DOCX instead.'}), 500
            
    except Exception as e:
        logger.error(f"PDF generation error: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to generate PDF: {str(e)}'}), 500

@bp.route('/ask_question', methods=['POST'])
@student_required
def ask_lesson_question():
    """
    Student Q&A on a lesson (Approach 3: stateless flag).

    Request body: lesson_id, question, allow_rag (optional, default False).
    When allow_rag=False and lesson doesn't cover the question but has a PDF,
    returns status=needs_rag_confirmation. Frontend shows Yes/No; on Yes,
    re-send the same question with allow_rag=true.
    """
    data = request.get_json() or {}
    lesson_id = data.get('lesson_id')
    question = (data.get('question') or '').strip()
    allow_rag = data.get('allow_rag', False)
    if not lesson_id or not question:
        return jsonify({'error': 'lesson_id and question are required'}), 400

    from app.models.models import LessonChatHistory

    user_id = session['user_id']
    allowed, retry_after = _check_and_record_lesson_qa_rate(int(user_id))
    if not allowed:
        return jsonify({
            'error': f'Too many requests in a short time. Please wait {retry_after} seconds and try again.',
            'code': 'USER_RATE_LIMITED',
            'retry_after': retry_after,
        }), 429

    user_lock = _get_lesson_qa_user_lock(int(user_id))
    acquired = user_lock.acquire(blocking=True, timeout=120)
    if not acquired:
        return jsonify({
            'error': 'Your previous message is still being processed. Please try again in a moment.',
            'code': 'CONCURRENT_REQUEST_TIMEOUT',
        }), 429

    try:
        result = invoke_lesson_qa(
            lesson_id=int(lesson_id),
            question=question,
            user_id=int(user_id),
            allow_rag=bool(allow_rag),
        )

        # Frontend shows Yes/No; on Yes, re-send with allow_rag=true
        if result.get('status') == 'needs_rag_confirmation':
            return jsonify({
                'status': 'needs_rag_confirmation',
                'permission_request': result.get('permission_request'),
            })

        answer = (result.get('answer') or '').strip()
        if not answer:
            return jsonify({'error': 'Failed to generate an answer.'}), 500

        # Persist the answer immediately; canonicalization runs asynchronously.
        chat_history_id = LessonChatHistory.save_qa(
            int(lesson_id),
            user_id,
            question,
            answer,
            canonical_question=question,
        )

        api_key = session.get('groq_api_key', '')
        app_obj = current_app._get_current_object()
        bg_thread = threading.Thread(
            target=_canonicalize_and_log_in_background,
            args=(app_obj, int(lesson_id), int(user_id), question, int(chat_history_id), api_key),
            daemon=True,
        )
        bg_thread.start()

        return jsonify({
            'status': 'completed',
            'answer': answer,
            # Return the original immediately; canonical value is updated in DB asynchronously.
            'canonical_question': question,
            'source': 'lesson_or_rag_graph',
            'confidence': 0.9,
        })
    except GroqRateLimitError as rl_exc:
        logger.warning(
            "ask_lesson_question: Groq rate limit hit for user %s — %s retry_after=%ds",
            user_id, rl_exc.info.kind, rl_exc.info.retry_after,
        )
        return jsonify({
            'error': 'The AI service is temporarily rate limited. Please wait and try again.',
            'code': rl_exc.info.kind,
            'retry_after': rl_exc.info.retry_after,
        }), 429
    except GroqBusyError as busy_exc:
        logger.warning("ask_lesson_question: Groq semaphore busy for user %s — %s", user_id, busy_exc)
        return jsonify({
            'error': 'The AI service is temporarily at capacity. Please try again in a moment.',
            'code': 'SERVICE_AT_CAPACITY',
            'retry_after': 10,
        }), 503
    except Exception as exc:
        logger.exception("ask_lesson_question: unexpected error for user %s lesson %s", user_id, lesson_id)
        return jsonify({
            'error': 'Failed to generate an answer. Please try again.',
            'code': 'INTERNAL_ERROR',
        }), 500
    finally:
        user_lock.release()


@bp.route('/lesson_chat_history/<int:lesson_id>', methods=['GET'])
@login_required
def get_lesson_chat_history(lesson_id):
    """Get chat history for a specific lesson"""
    try:
        from app.models.models import LessonChatHistory
        user_id = session['user_id']
        history = LessonChatHistory.get_lesson_chat_history(lesson_id, user_id)
        return jsonify({'history': history})
    except Exception as e:
        logger.error(f"Error getting lesson chat history: {str(e)}")
        return jsonify({'error': 'Failed to get lesson chat history'}), 500

@bp.route('/teacher/lesson_chat_history/<int:lesson_id>', methods=['GET'])
@teacher_required
def get_teacher_lesson_chat_history(lesson_id):
    """Get full chat history for a lesson for the owning teacher (all students for that lesson)."""
    try:
        # Ensure the requesting teacher owns this lesson (or is admin via decorator/role)
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404

        user_id = session['user_id']
        user_role = session.get('role', 'student')
        lesson_teacher_id = lesson.get('teacher_id')

        # Only the owning teacher (or admin) may see full lesson history
        if user_role != 'admin' and lesson_teacher_id != user_id:
            return jsonify({'error': 'Access denied'}), 403

        # Query all history rows for this lesson across users
        from app.models.database_models import LessonChatHistory as DBLessonChatHistory, User as DBUser
        db = get_db()
        rows = (
            db.query(DBLessonChatHistory, DBUser)
            .outerjoin(DBUser, DBLessonChatHistory.user_id == DBUser.id)
            .filter(DBLessonChatHistory.lesson_id == lesson_id)
            .order_by(DBLessonChatHistory.created_at.asc())
            .all()
        )

        history = []
        for record, user in rows:
            history.append({
                'question': record.question,
                'answer': record.answer,
                'created_at': record.created_at.isoformat() if record.created_at else None,
                'user_id': record.user_id,
                'user_email': getattr(user, 'email', None) if user else None
            })

        return jsonify({'history': history})
    except Exception as e:
        logger.error(f"Error getting teacher lesson chat history: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to get lesson chat history'}), 500

@bp.route('/clear_lesson_chat_history/<int:lesson_id>', methods=['DELETE'])
@login_required
def clear_lesson_chat_history(lesson_id):
    """Clear chat history for a specific lesson"""
    try:
        from app.models.models import LessonChatHistory
        user_id = session['user_id']
        LessonChatHistory.clear_lesson_chat_history(lesson_id, user_id)
        return jsonify({'message': 'Lesson chat history cleared successfully'})
    except Exception as e:
        logger.error(f"Error clearing lesson chat history: {str(e)}")
        return jsonify({'error': 'Failed to clear lesson chat history'}), 500

@bp.route('/lesson/<int:lesson_id>/create_version', methods=['POST'])
@teacher_required
def create_lesson_version(lesson_id):
    """Create a new version of an existing lesson"""
    try:
        # Check if original lesson exists and belongs to the teacher
        original_lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not original_lesson:
            return jsonify({'error': 'Original lesson not found'}), 404
        
        if original_lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        # Keep source_version_id as the lesson we're branching from (could be original or a child)
        source_version_id = lesson_id
        
        # For display defaults, load the root/original lesson's data if needed
        root_lesson_id = original_lesson.get('parent_lesson_id') or lesson_id
        original_lesson_data = LessonModel.get_lesson_by_id(root_lesson_id) or original_lesson
        
        data = request.get_json()
        
        # Debug logging
        logger.info(f"DEBUG: create_lesson_version - received data: {data}")
        logger.info(f"DEBUG: create_lesson_version - content field: {data.get('content', 'NOT_PROVIDED')}")
        logger.info(f"DEBUG: create_lesson_version - content length: {len(data.get('content', ''))}")
        
        # Check if new title already exists for this teacher (if title is being changed)
        new_title = data.get('title', original_lesson_data['title'])
        if new_title != original_lesson_data['title'] and LessonModel.check_title_exists(session['user_id'], new_title):
            return jsonify({'error': 'This lesson title is already used. Please choose a different title.'}), 400
        
        # Create new version: pass source_version_id; model will attach to root and flag the source
        new_lesson_id = LessonModel.create_new_version(
            original_lesson_id=source_version_id,
            teacher_id=session['user_id'],
            title=data.get('title', original_lesson_data['title']),
            summary=data.get('summary', original_lesson_data['summary']),
            learning_objectives=data.get('learning_objectives', original_lesson_data['learning_objectives']),
            focus_area=data.get('focus_area', original_lesson_data['focus_area']),
            grade_level=data.get('grade_level', original_lesson_data['grade_level']),
            content=data.get('content', original_lesson_data['content']),
            file_name=data.get('file_name', original_lesson_data.get('file_name'))
        )
        
        return jsonify({
            'success': True,
            'new_lesson_id': new_lesson_id,
            'message': 'New version created successfully'
        })
        
    except Exception as e:
        logger.error(f"Error creating lesson version: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to create lesson version: {str(e)}'}), 500

@bp.route('/check_title_exists', methods=['GET'])
@teacher_required
def check_title_exists():
    """Check if a lesson title already exists for the current teacher"""
    try:
        title = request.args.get('title', '').strip()
        if not title:
            return jsonify({'exists': False})
        
        teacher_id = session['user_id']
        exists = LessonModel.check_title_exists(teacher_id, title)
        
        return jsonify({'exists': exists})
        
    except Exception as e:
        logger.error(f"Error checking title existence: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to check title: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/create_ai_version', methods=['POST'])
@teacher_required
def create_ai_lesson_version(lesson_id):
    """Create an AI-improved version of an existing lesson"""
    try:
        # Check if original lesson exists and belongs to the teacher
        original_lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not original_lesson:
            return jsonify({'error': 'Original lesson not found'}), 404
        
        if original_lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        # Keep source_version_id as the lesson we're branching from (could be original or a child)
        source_version_id = lesson_id
        
        # For display defaults, load the root/original lesson's data if needed
        root_lesson_id = original_lesson.get('parent_lesson_id') or lesson_id
        original_lesson_data = LessonModel.get_lesson_by_id(root_lesson_id) or original_lesson
        
        data = request.get_json()
        improvement_prompt = data.get('improvement_prompt', '')
        
        # Use LessonService with central LLM provider (same as chat/RAG)
        # No need for per-user API key - uses central admin settings
        lesson_service = LessonService(api_key=None)
        
        # Generate improved lesson content
        improved_content = lesson_service.improve_lesson_content(
            lesson_id=lesson_id,
            current_content=original_lesson_data['content'],
            improvement_prompt=improvement_prompt
        )
        
        # Create new version with improved content, always using the original lesson ID
        new_lesson_id = LessonModel.create_new_version(
            original_lesson_id=lesson_id,
            teacher_id=session['user_id'],
            title=original_lesson_data['title'],
            summary=original_lesson_data['summary'],
            learning_objectives=original_lesson_data['learning_objectives'],
            focus_area=original_lesson_data['focus_area'],
            grade_level=original_lesson_data['grade_level'],
            content=improved_content,
            file_name=original_lesson_data.get('file_name')
        )
        
        return jsonify({
            'success': True,
            'new_lesson_id': new_lesson_id,
            'message': 'AI-improved version created successfully',
            'improved_lesson': {
                'title': original_lesson_data['title'],
                'summary': original_lesson_data['summary'],
                'learning_objectives': original_lesson_data['learning_objectives'],
                'focus_area': original_lesson_data['focus_area'],
                'grade_level': original_lesson_data['grade_level'],
                'content': improved_content
            }
        })
        
    except Exception as e:
        logger.error(f"Error creating AI lesson version: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to create AI lesson version: {str(e)}'}), 500

@bp.route('/faqs/<int:lesson_id>', methods=['GET'])
@teacher_required
def get_lesson_faqs(lesson_id):
    try:
        # Get lesson details first
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        db = get_db()
        canonical_expr = func.coalesce(DBLessonFAQ.canonical_question, DBLessonFAQ.question)
        faq_rows = (
            db.query(
                canonical_expr.label('question'),
                func.sum(DBLessonFAQ.count).label('count')
            )
            .filter(DBLessonFAQ.lesson_id == lesson_id)
            .group_by(canonical_expr)
            .order_by(desc('count'))
            .all()
        )
        
        # Load latest chat answers once, then map to normalized question keys.
        answer_rows = (
            db.query(
                DBLessonChatHistory.question,
                DBLessonChatHistory.canonical_question,
                DBLessonChatHistory.answer,
                DBLessonChatHistory.created_at,
                DBLessonChatHistory.id,
            )
            .filter(DBLessonChatHistory.lesson_id == lesson_id)
            .order_by(DBLessonChatHistory.created_at.desc(), DBLessonChatHistory.id.desc())
            .all()
        )
        latest_answer_by_key = {}
        for ans_row in answer_rows:
            row_key = _normalize_faq_question_text(
                ans_row.canonical_question or ans_row.question or ''
            )
            if not row_key or row_key in latest_answer_by_key:
                continue
            latest_answer_by_key[row_key] = ans_row.answer

        # Merge semantically equivalent FAQ rows into one canonical entry.
        merged = {}
        for row in faq_rows:
            raw_question = row[0] or ''
            normalized_key = _normalize_faq_question_text(raw_question) or str(raw_question).strip().lower()
            if not normalized_key:
                continue
            entry = merged.get(normalized_key)
            row_count = int(row[1] or 0)
            if not entry:
                merged[normalized_key] = {
                    'question_key': normalized_key,
                    'count': row_count,
                }
            else:
                entry['count'] += row_count

        # Format the FAQs to match the expected structure
        faqs = []
        for merged_row in sorted(merged.values(), key=lambda item: item['count'], reverse=True):
            # Determine version display text
            version_text = ""
            if lesson.get('parent_lesson_id'):
                version_text = f"v{lesson.get('version', 1)}"
            else:
                version_text = "v1 (Original)"
            question_key = merged_row['question_key']
            latest_answer = latest_answer_by_key.get(question_key)

            faqs.append({
                'question': _faq_display_question(question_key),
                'count': merged_row['count'],
                'times_asked': merged_row['count'],  # For compatibility with frontend
                'lessonTitle': f"{lesson.get('title', '')} {version_text}",
                'subject': lesson.get('focus_area', ''),
                'grade': lesson.get('grade_level', ''),
                'time_ago': 'Recently',  # Placeholder since timestamps aren't tracked
                'version': lesson.get('version', 1),
                'is_version': lesson.get('parent_lesson_id') is not None,
                'parent_lesson_id': lesson.get('parent_lesson_id'),
                'answer': latest_answer,
                'question_key': question_key,
            })
        
        return jsonify({'faqs': faqs})
        
    except Exception as e:
        logger.error(f"Error getting lesson FAQs: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to load lesson FAQs'}), 500

@bp.route('/faqs_count/<int:lesson_id>', methods=['GET'])
@login_required
def get_lesson_faq_count(lesson_id):
    try:
        db = get_db()
        total = (
            db.query(func.coalesce(func.sum(DBLessonFAQ.count), 0))
            .filter(DBLessonFAQ.lesson_id == lesson_id)
            .scalar()
            or 0
        )
        unique_qs = (
            db.query(DBLessonFAQ.id)
            .filter(DBLessonFAQ.lesson_id == lesson_id)
            .count()
        )
        return jsonify({'total_count': int(total), 'unique_count': int(unique_qs)})
    except Exception:
        return jsonify({'total_count': 0, 'unique_count': 0})

@bp.route('/faq_dashboard', methods=['GET'])
@teacher_required
def faq_dashboard():
    try:
        user_id = session['user_id']
        # Get all lessons for this teacher
        lessons = LessonModel.get_lessons_by_teacher(user_id)
        top_questions = []
        total_questions = 0
        recent_questions = []
        db = get_db()
        canonical_expr = func.coalesce(DBLessonFAQ.canonical_question, DBLessonFAQ.question)
        # For each lesson, get top questions
        for lesson in lessons:
            rows = (
                db.query(
                    canonical_expr.label('question'),
                    func.sum(DBLessonFAQ.count).label('count')
                )
                .filter(DBLessonFAQ.lesson_id == lesson['id'])
                .group_by(canonical_expr)
                .order_by(desc('count'))
                .limit(3)
                .all()
            )
            faqs = [{'question': row[0], 'count': int(row[1] or 0)} for row in rows]
            total_questions += sum(row['count'] for row in faqs)
            if faqs:
                top_questions.append({
                    'title': lesson['title'],
                    'subject': lesson.get('focus_area', ''),
                    'questions': faqs
                })
        # Get recent questions (last 24h) -- placeholder, as timestamps are not tracked
        # If you want real recent questions, you need to log timestamps in lesson_faq
        # For now, just return the most asked question per lesson
        for lesson in lessons:
            row = (
                db.query(
                    canonical_expr.label('question'),
                    func.sum(DBLessonFAQ.count).label('count')
                )
                .filter(DBLessonFAQ.lesson_id == lesson['id'])
                .group_by(canonical_expr)
                .order_by(desc('count'))
                .first()
            )
            if row:
                recent_questions.append({
                    'question': row[0],
                    'student_name': '',  # Not tracked
                    'lesson_title': lesson['title'],
                    'time_ago': 'Recently'
                })
        return jsonify({
            'total_questions': total_questions,
            'weekly_questions': total_questions,  # Placeholder
            'active_students': 0,  # Placeholder
            'top_questions': top_questions,
            'recent_questions': recent_questions
        })
    except Exception as e:
        logger.error(f"Error loading FAQ dashboard: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to load FAQ dashboard'}), 500 

@bp.route('/export_faq_dashboard', methods=['GET'])
@teacher_required
def export_faq_dashboard():
    try:
        user_id = session['user_id']
        lessons = LessonModel.get_lessons_by_teacher(user_id)
        from io import BytesIO
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        db = get_db()
        canonical_expr = func.coalesce(DBLessonFAQ.canonical_question, DBLessonFAQ.question)
        # Prepare data for export
        rows = []
        for lesson in lessons:
            faqs = (
                db.query(
                    canonical_expr.label('question'),
                    func.sum(DBLessonFAQ.count).label('count')
                )
                .filter(DBLessonFAQ.lesson_id == lesson['id'])
                .group_by(canonical_expr)
                .order_by(desc('count'))
                .all()
            )
            for faq in faqs:
                rows.append({
                    'lesson_title': lesson['title'],
                    'subject': lesson.get('focus_area', ''),
                    'question': faq[0],
                    'count': faq[1]
                })

        # Create Word document
        doc = Document()
        doc.add_heading('FAQ Dashboard Export', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Lesson Title'
        hdr_cells[1].text = 'Subject'
        hdr_cells[2].text = 'Question'
        hdr_cells[3].text = 'Count'
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                paragraph.runs[0].font.bold = True
                paragraph.runs[0].font.size = Pt(11)
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = str(row['lesson_title'])
            cells[1].text = str(row['subject'])
            cells[2].text = str(row['question'])
            cells[3].text = str(row['count'])
        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        from flask import make_response
        response = make_response(doc_io.getvalue())
        response.headers["Content-Disposition"] = "attachment; filename=faq_dashboard_export.docx"
        response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return response
    except Exception as e:
        logger.error(f"Error exporting FAQ dashboard: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to export FAQ dashboard'}), 500

@bp.route('/lesson/<int:lesson_id>/save_draft', methods=['POST'])
@teacher_required
def save_lesson_draft(lesson_id):
    """Save draft content for a lesson"""
    try:
        # Check if lesson exists and belongs to the teacher
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        data = request.get_json()
        draft_content = data.get('draft_content', '')
        
        success = LessonModel.save_draft_content(lesson_id, draft_content)
        
        if success:
            return jsonify({
                'success': True,
                'message': 'Draft content saved successfully'
            })
        else:
            return jsonify({'error': 'Failed to save draft content'}), 500
            
    except Exception as e:
        logger.error(f"Error saving draft content: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to save draft content: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/get_draft', methods=['GET'])
@teacher_required
def get_lesson_draft(lesson_id):
    """Get draft content for a lesson"""
    try:
        # Check if lesson exists and belongs to the teacher
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        draft_content = LessonModel.get_draft_content(lesson_id)
        
        # Get the current content from the lesson (this is what we update when finalizing)
        current_content = lesson.get('content') or lesson.get('original_content') or ''
        
        # Get the original content from the first version of this lesson
        original_content = lesson.get('original_content', '')
        if lesson.get('lesson_id'):
            # Get the first version (version_number = 1) to get true original content using ORM
            db = get_db()
            first_version = db.query(DBLesson).filter(
                DBLesson.lesson_id == lesson['lesson_id'],
                DBLesson.version_number == 1
            ).first()
            if first_version:
                # Use content if original_content is empty, otherwise use original_content
                original_content = first_version.original_content or first_version.content or ''
        
        # Log for debugging
        logger.info(f"get_draft for lesson {lesson_id}: current_content length={len(current_content)}, original_content length={len(original_content)}")
        
        return jsonify({
            'success': True,
            'draft_content': draft_content,
            'current_content': current_content,  # Current version's content
            'original_content': original_content   # True original content from first version
        })
            
    except Exception as e:
        logger.error(f"Error getting draft content: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to get draft content: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/review_by_ai', methods=['POST'])
@teacher_required
def review_by_ai(lesson_id):
    """Review lesson content by AI based on user query - provides fresh response without showing previous content"""
    try:
        # Check if lesson exists and belongs to the teacher
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        data = request.get_json()
        user_query = data.get('query', '')
        
        if not user_query.strip():
            return jsonify({'error': 'Query is required'}), 400
        
        # Get API key from session
        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured. Please set your API key first.'}), 400
        
        # Use LessonService to provide AI review
        lesson_service = LessonService(api_key=api_key)
        ai_response = lesson_service.review_by_ai(
            lesson_id=lesson_id,
            user_query=user_query
        )
        
        return jsonify({
            'success': True,
            'ai_response': ai_response,
            'message': 'AI review completed successfully'
        })
            
    except Exception as e:
        logger.error(f"Error getting AI review: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to get AI review: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/apply_prompt', methods=['POST'])
@teacher_required
def apply_prompt_to_lesson(lesson_id):
    """Apply AI prompt to lesson content and save as draft"""
    try:
        # Check if lesson exists and belongs to the teacher
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        data = request.get_json()
        prompt = data.get('prompt', '')
        
        if not prompt.strip():
            return jsonify({'error': 'Prompt is required'}), 400
        
        # Get current content (use draft if available, otherwise original)
        current_draft = LessonModel.get_draft_content(lesson_id)
        content_to_edit = current_draft if current_draft else lesson.get('original_content', lesson['content'])
        
        # Use LessonService with API key from session (user profile) or env - same source as RAG
        api_key = session.get('groq_api_key') or os.getenv('GROQ_API_KEY') or os.getenv('OPENAI_API_KEY')
        lesson_service = LessonService(api_key=api_key)
        improved_content = lesson_service.improve_lesson_content(
            lesson_id=lesson_id,
            current_content=content_to_edit,
            improvement_prompt=prompt
        )
        
        # Check if the improved content is a JSON string and extract the actual content
        import json
        
        # Recursive function to extract text content from JSON
        def extract_content(data):
            if isinstance(data, str):
                # Try to parse as JSON
                try:
                    parsed = json.loads(data)
                    return extract_content(parsed)
                except (json.JSONDecodeError, TypeError):
                    return data
            elif isinstance(data, dict):
                # Look for content in various fields
                for key in ['answer', 'response', 'content', 'text', 'message']:
                    if key in data:
                        return extract_content(data[key])
                # If no content field found, return string representation
                return str(data)
            elif isinstance(data, list):
                # Join list items
                return '\n'.join([extract_content(item) for item in data])
            else:
                return str(data)
        
        improved_content = extract_content(improved_content)
        
        # Save the improved content as draft
        success = LessonModel.save_draft_content(lesson_id, improved_content)
        
        if success:
            return jsonify({
                'success': True,
                'draft_content': improved_content,
                'message': 'Prompt applied successfully'
            })
        else:
            return jsonify({'error': 'Failed to save draft content'}), 500
            
    except Exception as e:
        logger.error(f"Error applying prompt: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to apply prompt: {str(e)}'}), 500

@bp.route('/lesson/<int:lesson_id>/finalize_version', methods=['POST'])
@teacher_required
def finalize_lesson_version(lesson_id):
    """Finalize draft content into a new lesson version"""
    try:
        # Check if lesson exists and belongs to the teacher
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        # Get draft content
        draft_content = LessonModel.get_draft_content(lesson_id)
        
        if not draft_content.strip():
            return jsonify({'error': 'No draft content to finalize. Apply a prompt first.'}), 400
        
        # Extract actual text content from JSON if needed
        import json
        
        def extract_content(data):
            if isinstance(data, str):
                try:
                    parsed = json.loads(data)
                    return extract_content(parsed)
                except (json.JSONDecodeError, TypeError):
                    return data
            elif isinstance(data, dict):
                for key in ['answer', 'response', 'content', 'text', 'message']:
                    if key in data:
                        return extract_content(data[key])
                return str(data)
            elif isinstance(data, list):
                return '\n'.join([extract_content(item) for item in data])
            else:
                return str(data)
        
        # Extract the actual content from JSON if the draft is in JSON format
        draft_content = extract_content(draft_content)
        
        # Determine the actual original lesson ID
        original_lesson_id = lesson.get('parent_lesson_id') or lesson_id
        
        # Create new version using draft content: pass source_version_id; model will attach to root and flag the source
        new_lesson_id = LessonModel.create_new_version_from_draft(
            original_lesson_id=lesson_id,
            teacher_id=session['user_id'],
            title=lesson['title'],
            summary=lesson['summary'],
            learning_objectives=lesson['learning_objectives'],
            focus_area=lesson['focus_area'],
            grade_level=lesson['grade_level'],
            draft_content=draft_content,
            file_name=lesson.get('file_name'),
            is_public=lesson.get('is_public', True)
        )
        
        # Update the new lesson with the final content in database
        new_lesson_model = LessonModel(new_lesson_id)
        new_lesson_model.update_lesson(content=draft_content)
        # Copy RAG thread from the lesson being finalized so "Ask Question" can use the uploaded PDF
        if lesson.get('rag_thread_id'):
            new_lesson_model.update_lesson(rag_thread_id=lesson['rag_thread_id'])
            logger.info("Copied rag_thread_id %s to finalized lesson %s", lesson['rag_thread_id'], new_lesson_id)
        
        # Delete FAISS index after storing in database
        try:
            lesson_service = LessonService(api_key=None)
            lesson_service._delete_faiss_index(new_lesson_id)
            logger.info(f"Deleted FAISS index for finalized lesson {new_lesson_id}")
        except Exception as e:
            logger.warning(f"Failed to delete FAISS index for lesson {new_lesson_id}: {str(e)}")
        
        # Clear the draft content from the current lesson
        LessonModel.clear_draft_content(lesson_id)
        
        return jsonify({
            'success': True,
            'new_lesson_id': new_lesson_id,
            'message': 'New version created successfully'
        })
        
    except Exception as e:
        logger.error(f"Error finalizing version: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to finalize version: {str(e)}'}), 500


#interactive chat with lesson

# ROUTE HANDLER
@bp.route('/lesson/<int:lesson_id>/interactive_chat', methods=['POST'])
@teacher_required
def interactive_chat(lesson_id):
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403

        data = request.get_json()
        user_query = data.get('query', '')

        if not user_query.strip():
            return jsonify({'error': 'Query is required'}), 400

        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured'}), 400

        # Get lesson data to pass form context
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404

        lesson_service = LessonService(api_key=api_key)
        
        # Check if document was uploaded (by checking if vector store exists)
        # For now, we'll assume document was uploaded if focus_area exists
        document_uploaded = bool(lesson.get('focus_area'))
        
        result = lesson_service.interactive_chat(
            lesson_id=lesson_id,
            user_query=user_query,
            subject=lesson.get('focus_area'),
            grade_level=lesson.get('grade_level'),
            focus_area=lesson.get('focus_area'),
            document_uploaded=document_uploaded,
            document_filename=None  # Can be enhanced to store filename
        )

        # If complete lesson is generated, save it to database
        if result.complete_lesson == "yes":
            # Update lesson content in database
            lesson_model = LessonModel(lesson_id)
            lesson_model.update_lesson(content=result.ai_response)
            logger.info(f"Complete lesson saved to database for lesson_id: {lesson_id}")

        # Return response with lesson update status
        return jsonify({
            'success': True,
            'ai_response': result.ai_response,
            'complete_lesson': result.complete_lesson,
            'lesson_id': lesson_id
        })

    except GroqRateLimitError as rl_exc:
        logger.warning(
            "interactive_chat: Groq rate limit for lesson %s — %s retry_after=%ds",
            lesson_id, rl_exc.info.kind, rl_exc.info.retry_after,
        )
        return jsonify({
            'error': 'The AI service is temporarily rate limited. Please wait and try again.',
            'code': rl_exc.info.kind,
            'retry_after': rl_exc.info.retry_after,
        }), 429
    except GroqBusyError as busy_exc:
        logger.warning("interactive_chat: Groq semaphore busy for lesson %s — %s", lesson_id, busy_exc)
        return jsonify({
            'error': 'The AI service is temporarily at capacity. Please try again in a moment.',
            'code': 'SERVICE_AT_CAPACITY',
            'retry_after': 10,
        }), 503
    except Exception as e:
        logger.error("Error in interactive chat lesson %s: %s", lesson_id, e, exc_info=True)
        return jsonify({'error': 'Failed to process chat. Please try again.', 'code': 'INTERNAL_ERROR'}), 500

@bp.route('/lesson/<int:lesson_id>/interactive_chat_stream', methods=['POST'])
@teacher_required
def interactive_chat_stream(lesson_id):
    """Interactive chat with streaming response using Server-Sent Events (SSE)"""
    try:
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403

        data = request.get_json()
        user_query = data.get('query', '')

        if not user_query.strip():
            return jsonify({'error': 'Query is required'}), 400

        api_key = session.get('groq_api_key')
        if not api_key:
            return jsonify({'error': 'API key not configured'}), 400

        # Get lesson data to pass form context
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404

        lesson_service = LessonService(api_key=api_key)
        
        # Check if document was uploaded
        document_uploaded = bool(lesson.get('focus_area'))
        
        # Import Response for streaming
        from flask import Response
        import json
        
        def generate():
            """Generator function for Server-Sent Events"""
            try:
                full_response = ""
                complete_lesson_status = "no"
                
                # Stream the response
                for chunk_text, is_complete, complete_lesson in lesson_service.interactive_chat_stream(
                    lesson_id=lesson_id,
                    user_query=user_query,
                    subject=lesson.get('focus_area'),
                    grade_level=lesson.get('grade_level'),
                    focus_area=lesson.get('focus_area'),
                    document_uploaded=document_uploaded,
                    document_filename=None
                ):
                    # Always accumulate the chunk text (even if empty for final message)
                    if chunk_text:
                        full_response += chunk_text
                        # Send chunk as SSE
                        yield f"data: {json.dumps({'chunk': chunk_text, 'is_complete': False, 'complete_lesson': 'no'})}\n\n"
                    
                    # Check if this is the final message
                    if is_complete:
                        complete_lesson_status = complete_lesson
                        # Send completion message with full response
                        yield f"data: {json.dumps({'chunk': '', 'is_complete': True, 'complete_lesson': complete_lesson_status, 'full_response': full_response})}\n\n"
                        
                        # If complete lesson is generated, save it to database
                        if complete_lesson_status == "yes" and full_response:
                            try:
                                lesson_model = LessonModel(lesson_id)
                                lesson_model.update_lesson(content=full_response)
                                logger.info(f"Complete lesson saved to database for lesson_id: {lesson_id}")
                            except Exception as e:
                                logger.error(f"Error saving lesson: {str(e)}")
                        
                        break
                
            except GroqRateLimitError as _rl:
                logger.warning("Streaming chat rate limit for lesson %s: %s", lesson_id, _rl.info.kind)
                error_data = json.dumps({
                    'error': 'The AI service is temporarily rate limited. Please wait and try again.',
                    'code': _rl.info.kind,
                    'retry_after': _rl.info.retry_after,
                    'is_complete': True,
                    'complete_lesson': 'no',
                })
                yield f"data: {error_data}\n\n"
            except GroqBusyError:
                logger.warning("Streaming chat: Groq semaphore busy for lesson %s", lesson_id)
                error_data = json.dumps({
                    'error': 'The AI service is temporarily at capacity. Please try again in a moment.',
                    'code': 'SERVICE_AT_CAPACITY',
                    'retry_after': 10,
                    'is_complete': True,
                    'complete_lesson': 'no',
                })
                yield f"data: {error_data}\n\n"
            except Exception as e:
                logger.error("Error in streaming chat lesson %s: %s", lesson_id, e, exc_info=True)
                error_data = json.dumps({
                    'error': 'Failed to generate response. Please try again.',
                    'code': 'INTERNAL_ERROR',
                    'is_complete': True,
                    'complete_lesson': 'no',
                })
                yield f"data: {error_data}\n\n"
        
        # Return streaming response with SSE headers
        return Response(
            generate(),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no',  # Disable buffering in nginx
                'Connection': 'keep-alive'
            }
        )

    except GroqRateLimitError as rl_exc:
        logger.warning("Streaming chat route rate limit for lesson %s: %s", lesson_id, rl_exc.info.kind)
        return jsonify({
            'error': 'The AI service is temporarily rate limited. Please wait and try again.',
            'code': rl_exc.info.kind,
            'retry_after': rl_exc.info.retry_after,
        }), 429
    except Exception as e:
        logger.error("Error in streaming chat route lesson %s: %s", lesson_id, e, exc_info=True)
        return jsonify({'error': 'Failed to start streaming. Please try again.', 'code': 'INTERNAL_ERROR'}), 500

@bp.route('/lesson/<int:lesson_id>/clear_draft', methods=['DELETE'])
@teacher_required
def clear_lesson_draft(lesson_id):
    """Clear draft content for a lesson"""
    try:
        # Check if lesson exists and belongs to the teacher
        lesson = LessonModel.get_lesson_by_id(lesson_id)
        if not lesson:
            return jsonify({'error': 'Lesson not found'}), 404
        
        if lesson['teacher_id'] != session['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        success = LessonModel.clear_draft_content(lesson_id)
        
        if success:
            return jsonify({
                'success': True,
                'message': 'Draft content cleared successfully'
            })
        else:
            return jsonify({'error': 'Failed to clear draft content'}), 500
            
    except Exception as e:
        logger.error(f"Error clearing draft content: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to clear draft content: {str(e)}'}), 500 


@bp.route('/chatbot', methods=['GET'])
@teacher_required
def chatbot():
    """Render the chatbot interface"""
    try:
        # Get user's lessons for selection
        user_id = session.get('user_id')
        lessons = LessonModel.get_lessons_by_teacher(user_id)
        
        return render_template('chatbot.html', lessons=lessons)
    except Exception as e:
        logger.error(f"Error rendering chatbot: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to render chatbot: {str(e)}'}), 500