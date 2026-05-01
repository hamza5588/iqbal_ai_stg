# app/routes/chat.py
from flask import Blueprint, redirect, request, session, jsonify, render_template, url_for, send_file
from app.services import ChatService, PromptService
from app.models.models import SurveyModel, LessonModel, ConversationModel
from app.services.conversation_summary_service import ConversationSummaryService
# from app.utils.decorators import login_required
from app.utils.auth import login_required
from app.utils.routes import get_default_route_by_role
from app.utils.decorators import teacher_required
import logging
from app.utils.db import get_db
import time
from functools import lru_cache
import os

import torch
from app.services.voice_service import transcribe_audio_file, synthesize_text_to_wav, normalize_language

import logging

logger = logging.getLogger(__name__)

# ✅ Whisper / PyTorch configuration (server-safe)
# Some CPU environments (e.g. certain staging/production hosts) can fail with
# "RuntimeError: could not create a primitive" when using oneDNN/MKLDNN for conv1d.
# Disabling MKLDNN and forcing CPU+float32 keeps behavior correct while avoiding
# those hardware-specific crashes.
torch.backends.mkldnn.enabled = False  # avoid oneDNN primitive creation issues

logger = logging.getLogger(__name__)
bp = Blueprint('chat', __name__)

# Simple cache for token status (user_id -> (timestamp, data))
_token_status_cache = {}
CACHE_TTL = 2  # Cache for 2 seconds

@bp.route('/health')
def health_check():
    """Health check endpoint for container orchestration"""
    return jsonify({'status': 'healthy'}), 200


@bp.route('/teacher-dashboard')
@login_required
@teacher_required
def teacher_dashboard():
    """Render teacher dashboard (teacher-only). Teachers use this page only, not chat.html. Template: teacher_dashboard.html; assets: /teacher-static/."""
    try:
        return render_template('teacher_dashboard.html')
    except Exception as e:
        logger.error(f"Error serving teacher dashboard: {str(e)}")
        return redirect(url_for('chat.index'))


@bp.route('/')
@login_required
def index():
    """
    Canonical authenticated landing.

    `chat.html` is legacy and must not be the default entrypoint for the new UI.
    We always redirect to the role dashboard.
    """
    return redirect(get_default_route_by_role(session.get('role')))


@bp.route('/student-dashboard')
@login_required
def student_dashboard():
    """
    Render the new student dashboard UI.

    Teachers/admins are redirected to their canonical dashboards.
    """
    role = session.get('role')
    if role == 'admin':
        return redirect('/admin/')
    if role == 'teacher':
        return redirect(url_for('chat.teacher_dashboard'))
    return render_template('student_dashboard.html')


def _render_legacy_chat():
    """
    Legacy UI renderer (`chat.html`).
    Kept for backward compatibility, but only reachable via explicit /legacy/* routes.
    """
    has_submitted_survey = False
    subscription_tier = 'free'
    try:
        survey_model = SurveyModel(session['user_id'])
        has_submitted_survey = survey_model.has_submitted_survey()
    except Exception as e:
        logger.warning(f"Survey check failed in legacy chat: {e}")

    try:
        from app.models.database_models import User as DBUser
        db = get_db()
        user = db.query(DBUser).filter(DBUser.id == session['user_id']).first()
        if user and getattr(user, 'subscription_tier', None):
            subscription_tier = user.subscription_tier
    except Exception as e:
        logger.warning(f"Subscription tier check failed in legacy chat: {e}")

    try:
        return render_template(
            'chat.html',
            has_submitted_survey=has_submitted_survey,
            subscription_tier=subscription_tier,
        )
    except Exception as e:
        logger.error(f"Error rendering legacy chat template: {str(e)}", exc_info=True)
        return render_template('chat.html', has_submitted_survey=False, subscription_tier='free')


@bp.route('/legacy/chat')
@login_required
def legacy_chat():
    """Explicit legacy route for `chat.html` (old UI)."""
    return _render_legacy_chat()

# Add these routes to chat.py

from app.services import PromptService

@bp.route('/get_prompt')
def get_prompt():
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401
        
    try:
        prompt_service = PromptService(session['user_id'])
        current_prompt = prompt_service.get_prompt()
        return jsonify({'prompt': current_prompt})
    except Exception as e:
        logger.error(f"Error retrieving prompt: {str(e)}")
        return jsonify({'error': 'Failed to retrieve prompt'}), 500

@bp.route('/update_prompt', methods=['POST'])
def update_prompt():
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401
        
    try:
        data = request.json
        new_prompt = data.get('prompt')
   
        
        if not new_prompt:
            return jsonify({'error': 'Prompt is required'}), 400
            
        prompt_service = PromptService(session['user_id'])
        prompt_service.update_prompt(new_prompt)
        
        return jsonify({
            'success': True,
            'message': 'Prompt updated successfully'
        })
        
    except Exception as e:
        logger.error(f"Error updating prompt: {str(e)}")
        return jsonify({'error': 'Failed to update prompt'}), 500

@bp.route('/chat', methods=['POST'])
@login_required
def chat():
    """Handle chat messages and generate responses"""
    try:
        data = request.json
        user_input = data.get('input', '').strip()
        conversation_id = data.get('conversation_id')

        if not user_input:
            return jsonify({'error': 'Empty message'}), 400

        # Check LLM provider and ensure API key is available when Groq is selected
        import os
        from app.utils.db import get_db
        from app.models.database_models import SystemSettings
        db = get_db()
        setting = db.query(SystemSettings).filter(SystemSettings.key == 'llm_provider').first()
        provider = setting.value if setting else os.getenv('LLM_PROVIDER', 'openai').lower()
        
        # When Groq is selected, API key is mandatory - no fallback to OpenAI
        api_key = session.get('groq_api_key', '')
        if provider == 'groq' and not api_key:
            # Try to get from database
            from app.models import UserModel
            user = UserModel.get_user_by_id(session['user_id'])
            api_key = user.get('groq_api_key', '') if user else ''
            
            if not api_key:
                return jsonify({
                    'error': 'Groq API key is required. Please configure your Groq API key using the key icon in the chat interface.',
                    'requires_api_key': True,
                    'provider': 'groq'
                }), 400
        
        # Initialize services
        chat_service = ChatService(session['user_id'], api_key)
        prompt_service = PromptService(session['user_id'])

        # Get system prompt
        system_prompt = prompt_service.get_prompt()

        # Process message and generate response
        try:
            result = chat_service.process_message(
                message=user_input,
                conversation_id=conversation_id
            )
            return jsonify(result)
        except ValueError as e:
            # Handle missing API key errors specifically
            error_msg = str(e)
            if 'API key' in error_msg or 'Groq' in error_msg:
                logger.error(f"API key error: {error_msg}")
                return jsonify({
                    'error': error_msg,
                    'requires_api_key': True,
                    'provider': provider
                }), 400
            raise
        except Exception as e:
            logger.error(f"Error generating response: {str(e)}")
            # Check if it's a Groq API key issue
            if provider == 'groq' and ('API key' in str(e) or 'Groq' in str(e)):
                return jsonify({
                    'error': 'Groq API key is required. Please configure your Groq API key using the key icon in the chat interface.',
                    'requires_api_key': True,
                    'provider': 'groq'
                }), 400
            return jsonify({
                'error': """1:Your free key has expired,please login after 24 hours
                            2:Create another gmail account and login
                            3:Login to paid service-avalible in oct 2025"""
            }), 500

    except Exception as e:
        logger.error(f"Chat error: {str(e)}")
        return jsonify({'error': 'An error occurred'}), 500

@bp.route('/create_conversation', methods=['POST'])
@login_required
def create_conversation():
    """Create a new conversation"""
    try:
        data = request.json
        title = data.get('title', 'New Conversation')
        
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        conversation_id = chat_service.create_conversation(title)
        
        return jsonify({
            'conversation_id': conversation_id,
            'title': title
        })
    except Exception as e:
        logger.error(f"Error creating conversation: {str(e)}")
        return jsonify({'error': 'Failed to create conversation'}), 500

@bp.route('/save_message', methods=['POST'])
@login_required
def save_message():
    """Persist a single message to DB conversation history."""
    try:
        data = request.get_json(silent=True) or {}
        raw_message = data.get('message', '')
        message = str(raw_message).strip()
        if not message:
            return jsonify({'error': 'Message is required'}), 400

        role = str(data.get('role', '')).strip().lower()
        if role == 'assistant':
            role = 'bot'
        if role not in ('user', 'bot'):
            return jsonify({'error': 'Invalid role. Use user or bot.'}), 400

        conv_model = ConversationModel(session['user_id'])
        conversation_id = data.get('conversation_id')
        title = str(data.get('title') or 'New Conversation').strip() or 'New Conversation'

        if conversation_id is None or str(conversation_id).strip() == '':
            conversation_id = conv_model.create_conversation(title[:120])
        else:
            try:
                conversation_id = int(conversation_id)
            except (TypeError, ValueError):
                return jsonify({'error': 'Invalid conversation_id'}), 400
            conversation = conv_model.get_conversation_by_id(conversation_id)
            if not conversation:
                return jsonify({'error': 'Conversation not found or access denied'}), 404

        message_id = conv_model.save_message(
            conversation_id=conversation_id,
            message=message,
            role=role,
        )
        return jsonify({
            'success': True,
            'conversation_id': conversation_id,
            'message_id': message_id,
        })
    except Exception as e:
        logger.error(f"Error saving message: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to save message'}), 500

@bp.route('/get_conversations')
@login_required
def get_conversations():
    """Get user's recent conversations"""
    try:
        limit = request.args.get('limit', type=int)
        if limit is None:
            limit = 200
        # Guardrails: prevent pathological values but keep room for long histories
        limit = max(1, min(limit, 1000))
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        conversations = chat_service.get_recent_conversations(limit=limit)
        return jsonify({'conversations': conversations})  # <-- wrap in dict for frontend
    except Exception as e:
        logger.error(f"Error retrieving conversations: {str(e)}")
        return jsonify({'error': 'Failed to retrieve conversations'}), 500

@bp.route('/get_messages/<int:conversation_id>')
@login_required
def get_messages(conversation_id):
    """Get messages for a specific conversation. Includes thread_id when this conversation has an uploaded PDF."""
    try:
        user_id = session['user_id']
        chat_service = ChatService(user_id, session['groq_api_key'])
        messages = chat_service.get_conversation_messages(conversation_id)
        # Resolve RAG thread_id and filename for this conversation (for correct chat context and preamble)
        thread_id = None
        uploaded_filename = None
        try:
            from app.models.database_models import RAGThread
            db = get_db()
            prefix = f"user_{user_id}_conv_{conversation_id}_"
            row = db.query(RAGThread.thread_id, RAGThread.filename).filter(
                RAGThread.user_id == user_id,
                RAGThread.thread_id.like(prefix + "%"),
                RAGThread.has_document == True,
            ).order_by(RAGThread.created_at.desc()).first()
            if row:
                thread_id = row[0] if isinstance(row, (tuple, list)) else row.thread_id
                uploaded_filename = (row[1] if isinstance(row, (tuple, list)) and len(row) > 1 else getattr(row, 'filename', None)) or None
        except Exception:
            pass
        return jsonify({
            'messages': messages,
            'thread_id': thread_id,
            'uploaded_filename': uploaded_filename,
        })
    except Exception as e:
        logger.error(f"Error retrieving messages: {str(e)}")
        return jsonify({'error': 'Failed to retrieve messages'}), 500

@bp.route('/get_conversation/<int:conversation_id>')
@login_required
def get_conversation(conversation_id):
    """Get conversation details including title"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        conversation = chat_service.get_conversation_details(conversation_id)
        if conversation:
            return jsonify(conversation)
        else:
            return jsonify({'error': 'Conversation not found'}), 404
    except Exception as e:
        logger.error(f"Error retrieving conversation: {str(e)}")
        return jsonify({'error': 'Failed to retrieve conversation'}), 500


@bp.route('/api/conversations/<int:conversation_id>/summary', methods=['GET'])
@login_required
def get_conversation_summary(conversation_id):
    """Return latest summary plus freshness metadata."""
    try:
        user_id = session['user_id']
        lesson_id_raw = request.args.get('lesson_id')
        lesson_id = None
        if lesson_id_raw not in (None, ''):
            try:
                lesson_id = int(lesson_id_raw)
            except (TypeError, ValueError):
                return jsonify({'error': 'Invalid lesson_id'}), 400

        conv = ConversationSummaryService.get_conversation_for_user(conversation_id, user_id)
        if not conv:
            return jsonify({'error': 'Conversation not found or access denied'}), 404

        lesson_cutoff = ConversationSummaryService.get_lesson_cutoff(
            conversation_id=conversation_id,
            lesson_id=lesson_id,
            user_id=user_id,
        )
        latest = ConversationSummaryService.get_latest_summary(conversation_id, lesson_id=lesson_id)
        is_outdated = ConversationSummaryService.is_summary_outdated(
            conversation_id,
            latest,
            lesson_cutoff=lesson_cutoff,
        )
        if not latest:
            return jsonify({
                'success': True,
                'summary': None,
                'generated_at': None,
                'last_message_id': None,
                'last_message_timestamp': None,
                'is_outdated': True,
                'message': 'No summary available yet',
            })

        return jsonify({
            'success': True,
            'summary': latest.summary_text,
            'generated_at': latest.generated_at.isoformat() if latest.generated_at else None,
            'last_message_id': latest.last_message_id,
            'last_message_timestamp': latest.last_message_timestamp.isoformat() if latest.last_message_timestamp else None,
            'is_outdated': is_outdated,
        })
    except ValueError as e:
        return jsonify({'error': str(e)}), 404
    except Exception as e:
        logger.error(f"Error retrieving summary: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to retrieve summary'}), 500


@bp.route('/api/conversations/<int:conversation_id>/summary/regenerate', methods=['POST'])
@login_required
def regenerate_conversation_summary(conversation_id):
    """Generate and store latest summary."""
    try:
        user_id = session['user_id']
        data = request.get_json(silent=True) or {}
        lesson_id = data.get('lesson_id')
        if lesson_id in ('', None):
            lesson_id = None
        elif not isinstance(lesson_id, int):
            try:
                lesson_id = int(lesson_id)
            except (TypeError, ValueError):
                return jsonify({'error': 'Invalid lesson_id'}), 400

        force = bool(data.get('force', True))
        result = ConversationSummaryService.generate_and_persist_summary(
            conversation_id=conversation_id,
            user_id=user_id,
            lesson_id=lesson_id,
            force=force,
        )
        return jsonify({
            'success': True,
            'summary': result.get('summary'),
            'generated_at': result.get('generated_at').isoformat() if result.get('generated_at') else None,
            'last_message_id': result.get('last_message_id'),
            'last_message_timestamp': (
                result.get('last_message_timestamp').isoformat()
                if result.get('last_message_timestamp')
                else None
            ),
            'is_outdated': result.get('is_outdated', False),
            'was_regenerated': result.get('was_regenerated', False),
        })
    except ValueError as e:
        return jsonify({'error': str(e)}), 404
    except Exception as e:
        logger.error(f"Error regenerating summary: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to regenerate summary'}), 500

@bp.route('/delete_conversation/<int:conversation_id>', methods=['DELETE'])
@login_required
def delete_conversation(conversation_id):
    """Delete a conversation"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        chat_service.delete_conversation(conversation_id)
        return jsonify({'message': 'Conversation deleted successfully'})
    except Exception as e:
        logger.error(f"Error deleting conversation: {str(e)}")
        return jsonify({'error': 'Failed to delete conversation'}), 500

@bp.route('/delete_all_conversations', methods=['DELETE'])
@login_required
def delete_all_conversations():
    """Delete all conversations for the current user"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        chat_service.reset_all_conversations()
        return jsonify({'message': 'All conversations deleted successfully'})
    except Exception as e:
        logger.error(f"Error deleting all conversations: {str(e)}")
        return jsonify({'error': 'Failed to delete conversations'}), 500


@bp.route('/duplicate_conversation/<int:conversation_id>', methods=['POST'])
@login_required
def duplicate_conversation(conversation_id):
    """
    Duplicate a conversation and all of its messages for the current user.
    """
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        new_id = chat_service.duplicate_conversation(conversation_id)
        conversation = chat_service.get_conversation_details(new_id)
        return jsonify({
            'conversation_id': new_id,
            'conversation': conversation,
        })
    except ValueError as ve:
        logger.warning(f"Duplicate conversation failed: {ve}")
        return jsonify({'error': str(ve)}), 404
    except Exception as e:
        logger.error(f"Error duplicating conversation: {str(e)}")
        return jsonify({'error': 'Failed to duplicate conversation'}), 500


@bp.route('/update_conversation_title/<int:conversation_id>', methods=['PUT'])
@login_required
def update_conversation_title(conversation_id):
    """Update the title of a conversation"""
    try:
        data = request.json
        new_title = data.get('title', '').strip()
        
        if not new_title:
            return jsonify({'error': 'Title cannot be empty'}), 400
        
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        success = chat_service.update_conversation_title(conversation_id, new_title)
        
        if success:
            return jsonify({'message': 'Title updated successfully', 'title': new_title})
        else:
            return jsonify({'error': 'Conversation not found or access denied'}), 404
            
    except Exception as e:
        logger.error(f"Error updating conversation title: {str(e)}")
        return jsonify({'error': 'Failed to update title'}), 500

@bp.route('/download_chat/<int:conversation_id>')
@login_required
def download_chat(conversation_id):
    """Download a chat conversation as a Word document"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        messages = chat_service.get_conversation_messages(conversation_id)
        
        # Create a new Word document
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('Chat Conversation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add messages
        for msg in messages:
            role = "Mr. Potter" if msg['role'] == 'bot' else "User"
            p = doc.add_paragraph()
            p.add_run(f"{role}: ").bold = True
            p.add_run(msg['message'])
            p.paragraph_format.space_after = Pt(12)
        
        # Save the document to a BytesIO object
        from io import BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Create response with appropriate headers
        from flask import make_response
        response = make_response(doc_io.getvalue())
        response.headers["Content-Disposition"] = f"attachment; filename=chat_conversation_{conversation_id}.docx"
        response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        return response
    except Exception as e:
        logger.error(f"Error downloading chat: {str(e)}")
        return jsonify({'error': 'Failed to download chat'}), 500


@bp.route('/download_last_ai_message/<int:conversation_id>')
@login_required
def download_last_ai_message(conversation_id):
    """Download the last AI message from a chat as a Word document"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        messages = chat_service.get_conversation_messages(conversation_id)

        if not messages:
            return jsonify({'error': 'No messages found in this conversation'}), 404

        # Find the last AI/bot message
        last_bot_message = None
        for msg in reversed(messages):
            if msg.get('role') == 'bot' and msg.get('message'):
                last_bot_message = msg['message']
                break

        if not last_bot_message:
            return jsonify({'error': 'No AI messages found to download'}), 404

        # Create a new Word document containing only the last AI message
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        from flask import make_response

        doc = Document()

        # Add a title
        title = doc.add_heading('AI Lecture', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add some spacing
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)

        # Add the AI message content
        content_paragraph = doc.add_paragraph(last_bot_message)
        content_paragraph.paragraph_format.space_after = Pt(12)

        # Save document to buffer
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Prepare response
        response = make_response(doc_io.getvalue())
        response.headers["Content-Disposition"] = f"attachment; filename=ai_lecture_{conversation_id}.docx"
        response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return response
    except Exception as e:
        logger.error(f"Error downloading last AI message: {str(e)}")
        return jsonify({'error': 'Failed to download AI lecture'}), 500

@bp.route('/get_token_usage')
@login_required
def get_token_usage():
    """Get current token usage information"""
    try:
        chat_service = ChatService(session['user_id'], session['groq_api_key'])
        token_usage = chat_service.get_token_usage()
        return jsonify(token_usage)
    except Exception as e:
        logger.error(f"Error getting token usage: {str(e)}")
        return jsonify({
            'daily_limit': '100,000',
            'used_tokens': '0',
            'requested_tokens': '0',
            'wait_time': None
        }), 500
    



# TOKEN USAGE ROUTE 

@bp.route('/token_status', methods=['GET'])
@login_required
def get_token_status():
    try:
        user_id = session['user_id']
        current_api_key = session.get('groq_api_key')
        current_time = time.time()
        
        # Check cache first
        cache_key = f"{user_id}_{current_api_key}"
        if cache_key in _token_status_cache:
            cached_time, cached_data = _token_status_cache[cache_key]
            if current_time - cached_time < CACHE_TTL:
                # Return cached data
                return jsonify(cached_data)
        
        # Cache miss or expired - fetch fresh data
        # Try to get token usage directly from database first (faster)
        try:
            from app.utils.db import get_token_usage
            db_usage = get_token_usage(user_id)
            daily_limit = 100000  # Default limit
            
            used = db_usage['today']['tokens_used']
            remaining = max(0, daily_limit - used)
            
            response_data = {
                'used': used,
                'remaining': remaining,
                'limit': daily_limit,
                'reset_in': 0,  # Will be calculated by frontend if needed
                'api_key_changed': False
            }
            
            # Cache the response
            _token_status_cache[cache_key] = (current_time, response_data)
            
            # Clean old cache entries (keep only last 100)
            if len(_token_status_cache) > 100:
                # Remove oldest entries
                sorted_cache = sorted(_token_status_cache.items(), key=lambda x: x[1][0])
                for key, _ in sorted_cache[:-100]:
                    del _token_status_cache[key]
            
            return jsonify(response_data)
            
        except Exception as db_error:
            logger.warning(f"Direct DB query failed, falling back to ChatService: {str(db_error)}")
            # Fallback to ChatService if direct DB query fails
            chat_service = ChatService(user_id, current_api_key)
            
            # Verify the API key matches what's being used
            if chat_service.chat_model.api_key != current_api_key:
                chat_service.chat_model.api_key = current_api_key  # This will trigger reset
                
            token_status = chat_service.chat_model.get_token_status()
            
            response_data = {
                'used': token_status['used'],
                'remaining': token_status['remaining'],
                'limit': token_status['limit'],
                'reset_in': token_status['reset_in'],
                'api_key_changed': False  # Can be used by frontend if needed
            }
            
            # Cache the response
            _token_status_cache[cache_key] = (current_time, response_data)
            
            # Clean old cache entries (keep only last 100)
            if len(_token_status_cache) > 100:
                # Remove oldest entries
                sorted_cache = sorted(_token_status_cache.items(), key=lambda x: x[1][0])
                for key, _ in sorted_cache[:-100]:
                    del _token_status_cache[key]
            
            return jsonify(response_data)
            
    except Exception as e:
        logger.error(f"Error getting token status: {str(e)}", exc_info=True)
        # Return cached data if available, even if expired
        cache_key = f"{session.get('user_id')}_{session.get('groq_api_key')}"
        if cache_key in _token_status_cache:
            _, cached_data = _token_status_cache[cache_key]
            return jsonify(cached_data)
        # Return error response
        return jsonify({
            'error': 'Failed to get token status',
            'used': 0,
            'remaining': 0,
            'limit': 100000,
            'reset_in': 0
        }), 500

@bp.route('/user_info')
def get_user_info():
    """Get current user information including role"""
    if 'user_id' not in session:
        return jsonify({'error': 'Not authenticated'}), 401
    
    try:
        from app.models import UserModel
        user_model = UserModel(session['user_id'])
        user_info = UserModel.get_user_by_id(session['user_id'])
        
        if not user_info:
            return jsonify({'error': 'User not found'}), 404
        
        return jsonify({
            'success': True,
            'user': {
                'id': user_info['id'],
                'username': user_info['username'],
                'role': user_info.get('role', 'student'),
                'class_standard': user_info['class_standard'],
                'medium': user_info['medium']
            }
        })
    except Exception as e:
        logger.error(f"Error getting user info: {str(e)}")
        return jsonify({'error': 'Failed to get user info'}), 500

def _speech_to_text_impl():
    """
    Convert uploaded speech audio to text using local Whisper model.
    Supports optional language override: auto/en/ur/hi.
    """
    try:
        if 'audio' not in request.files:
            return jsonify({'error': 'No audio file provided', 'code': 'AUDIO_REQUIRED'}), 400

        audio_file = request.files['audio']
        if audio_file.filename == '':
            return jsonify({'error': 'Empty audio filename', 'code': 'AUDIO_REQUIRED'}), 400

        from tempfile import NamedTemporaryFile
        lang = normalize_language(request.form.get("language"), default="auto")
        tmp_path = None
        try:
            with NamedTemporaryFile(delete=False, suffix=".webm") as tmp:
                audio_file.save(tmp.name)
                tmp_path = tmp.name

            result, error = transcribe_audio_file(tmp_path, language=lang)
            if error:
                status = 422 if error.get("code") == "NO_SPEECH" else 400
                if error.get("code") in {"STT_FAILED"}:
                    status = 500
                return jsonify({'error': error.get("message"), 'code': error.get("code")}), status

            return jsonify({
                'text': result.get("text"),
                'language': result.get("language"),
                'meta': {
                    'avg_no_speech_prob': result.get("avg_no_speech_prob"),
                }
            })
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass

    except Exception as e:
        logger.error(f"Error in speech_to_text: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to transcribe audio', 'code': 'STT_EXCEPTION'}), 500


@bp.route('/speech-to-text', methods=['POST'])
@login_required
def speech_to_text():
    return _speech_to_text_impl()


@bp.route('/api/stt', methods=['POST'])
@login_required
def speech_to_text_legacy():
    return _speech_to_text_impl()

# @bp.route('/api/stt', methods=['POST'])
# @login_required
# def speech_to_text():
#     """
#     Convert uploaded speech audio to text using OpenAI Whisper.

#     Expects multipart/form-data with field "audio".
#     Returns JSON: {"text": "..."} on success.
#     """
#     try:
#         if 'audio' not in request.files:
#             return jsonify({'error': 'No audio file provided'}), 400

#         audio_file = request.files['audio']
#         if audio_file.filename == '':
#             return jsonify({'error': 'Empty audio filename'}), 400

#         # OpenAI client for Whisper
#         client = _get_openai_client()

#         # Whisper works best with binary file-like objects
#         # We read into memory here as recordings are short (voice messages)
#         from tempfile import NamedTemporaryFile

#         with NamedTemporaryFile(delete=False, suffix=".webm") as tmp:
#             audio_file.save(tmp.name)
#             tmp_path = tmp.name

#         with open(tmp_path, "rb") as f:
#             transcription = client.audio.transcriptions.create(
#                 model="whisper-1",
#                 file=f,
#                 response_format="json"
#             )

#         text = getattr(transcription, "text", None) or transcription.get("text")  # handle both object/dict
#         if not text:
#             return jsonify({'error': 'Transcription failed'}), 500

#         return jsonify({'text': text})

#     except Exception as e:
#         logger.error(f"Error in speech_to_text: {str(e)}", exc_info=True)
#         return jsonify({'error': 'Failed to transcribe audio'}), 500

@bp.route('/chatbot', methods=['GET'])
@teacher_required
def chatbot():
    """Render the chatbot interface"""
    try:
        # Get user's lessons for selection
        user_id = session.get('user_id')
        lessons = LessonModel.get_lessons_by_teacher(user_id)
        
        return render_template('chatbot.html', lessons=lessons)
    except Exception as e:
        logger.error(f"Error rendering chatbot: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to render chatbot: {str(e)}'}), 500


@bp.route('/chatbot_update', methods=['GET'])
@login_required
def chatbot_update():
    """Render the PDF chat interface"""
    try:
        return render_template('chatbot_update.html')
    except Exception as e:
        logger.error(f"Error rendering chatbot_update page: {str(e)}", exc_info=True)
        return jsonify({'error': f'Failed to render page: {str(e)}'}), 500




def _text_to_speech_impl():
    try:
        data = request.get_json() or {}
        text = (data.get('text') or '').strip()
        language = normalize_language(data.get('language'), default='auto')

        if not text:
            return jsonify({'error': 'Text is required', 'code': 'TEXT_REQUIRED'}), 400

        audio_fp, meta_or_error = synthesize_text_to_wav(text=text, lang_hint=language)
        if audio_fp is None:
            code = meta_or_error.get("code") if isinstance(meta_or_error, dict) else "TTS_FAILED"
            message = meta_or_error.get("message") if isinstance(meta_or_error, dict) else "Failed to generate speech"
            status = 503 if code == "VOICE_MODEL_MISSING" else 500
            return jsonify({'error': message, 'code': code, 'meta': meta_or_error}), status

        audio_format = (meta_or_error or {}).get("format", "wav") if isinstance(meta_or_error, dict) else "wav"
        mimetype = 'audio/mpeg' if audio_format == "mp3" else 'audio/wav'
        download_name = 'tts.mp3' if audio_format == "mp3" else 'tts.wav'
        return send_file(
            audio_fp,
            mimetype=mimetype,
            as_attachment=False,
            download_name=download_name
        )
    except Exception as e:
        logger.error(f"Error in text_to_speech: {str(e)}", exc_info=True)
        return jsonify({'error': 'Failed to generate speech', 'code': 'TTS_EXCEPTION'}), 500


@bp.route('/text-to-speech', methods=['POST'])
@login_required
def text_to_speech():
    return _text_to_speech_impl()


@bp.route('/api/tts', methods=['POST'])
@login_required
def text_to_speech_legacy():
    return _text_to_speech_impl()



# def text_to_speech():
#     """
#     Convert text to speech using gTTS and return an audio stream.
#     This replaces browser-based SpeechSynthesis so it also works for Urdu and other languages.
#     """
#     try:
#         data = request.get_json() or {}
#         text = (data.get('text') or '').strip()

#         if not text:
#             return jsonify({'error': 'Text is required'}), 400

#         # Detect language; fallback to English on failure
#         try:
#             lang = detect(text)
#         except Exception as e:
#             logger.warning(f"Language detection failed, defaulting to 'en': {str(e)}")
#             lang = 'en'

#         # Generate speech with gTTS
#         tts = gTTS(text=text, lang=lang)
#         audio_fp = BytesIO()
#         tts.write_to_fp(audio_fp)
#         audio_fp.seek(0)

#         # Return audio as an MP3 stream
#         return send_file(
#             audio_fp,
#             mimetype='audio/mpeg',
#             as_attachment=False,
#             download_name='tts.mp3'
#         )
#     except Exception as e:
#         logger.error(f"Error in text_to_speech: {str(e)}", exc_info=True)
#         return jsonify({'error': 'Failed to generate speech'}), 500
